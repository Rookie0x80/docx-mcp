"""Document management operations for Word documents."""

import os
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches

from ..models.responses import OperationResponse
from ..models.tables import TableInfo
from ..utils.exceptions import (
    DocumentNotFoundError,
    DocumentAccessError,
    TableOperationError,
)
from ..utils.validation import validate_file_path


class DocumentManager:
    """Manages Word document operations."""
    
    def __init__(self):
        """Initialize document manager."""
        self._documents: Dict[str, Document] = {}
    
    def open_document(self, file_path: str, create_if_not_exists: bool = True) -> OperationResponse:
        """
        Open or create a Word document.
        
        Args:
            file_path: Path to the document file
            create_if_not_exists: Whether to create the document if it doesn't exist
            
        Returns:
            OperationResponse with status and message
        """
        try:
            path = Path(file_path)
            is_new = False
            
            if path.exists():
                # Validate and open existing file
                validate_file_path(file_path, must_exist=True)
                document = Document(str(path))
                message = f"Opened existing document: {file_path}"
            else:
                if create_if_not_exists:
                    # Create new document
                    document = Document()
                    is_new = True
                    message = f"Created new document: {file_path}"
                else:
                    raise DocumentNotFoundError(f"Document not found: {file_path}")
            
            # Cache the document
            self._documents[file_path] = document
            
            # Get document info
            table_count = len(document.tables)
            paragraph_count = len(document.paragraphs)
            
            data = {
                "file_path": file_path,
                "table_count": table_count,
                "paragraph_count": paragraph_count,
                "is_new": is_new
            }
            
            return OperationResponse.success(message, data)
            
        except Exception as e:
            return OperationResponse.error(f"Failed to open document: {str(e)}")
    
    def save_document(self, file_path: str, save_as: Optional[str] = None) -> OperationResponse:
        """
        Save a Word document.
        
        Args:
            file_path: Path to the document file
            save_as: Optional path to save as a different file
            
        Returns:
            OperationResponse with status and message
        """
        try:
            if file_path not in self._documents:
                return OperationResponse.error(f"Document not loaded: {file_path}")
            
            document = self._documents[file_path]
            save_path = save_as if save_as else file_path
            
            # Validate save path
            validate_file_path(save_path, must_exist=False)
            
            # Save the document
            document.save(save_path)
            
            # Update cache if saving with a new name
            if save_as:
                self._documents[save_as] = document
            
            message = f"Document saved to: {save_path}"
            data = {"file_path": save_path}
            
            return OperationResponse.success(message, data)
            
        except Exception as e:
            return OperationResponse.error(f"Failed to save document: {str(e)}")
    
    def get_document(self, file_path: str) -> Optional[Document]:
        """
        Get a loaded document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object if loaded, None otherwise
        """
        return self._documents.get(file_path)
    
    def get_or_load_document(self, file_path: str, create_if_not_exists: bool = False) -> Document:
        """
        Get a document, loading it automatically if not already loaded.
        
        Args:
            file_path: Path to the document file
            create_if_not_exists: Whether to create the document if it doesn't exist
            
        Returns:
            Document object
            
        Raises:
            DocumentNotFoundError: If document doesn't exist and create_if_not_exists is False
            DocumentAccessError: If there's an error accessing the document
        """
        # Check if document is already loaded
        if file_path in self._documents:
            return self._documents[file_path]
        
        # Load the document
        try:
            path = Path(file_path)
            
            if path.exists():
                # Validate and open existing file
                validate_file_path(file_path, must_exist=True)
                document = Document(str(path))
            else:
                if create_if_not_exists:
                    # Create new document
                    document = Document()
                else:
                    raise DocumentNotFoundError(f"Document not found: {file_path}")
            
            # Cache the document
            self._documents[file_path] = document
            
            return document
            
        except DocumentNotFoundError:
            raise
        except Exception as e:
            raise DocumentAccessError(f"Failed to load document {file_path}: {str(e)}")
    
    def close_document(self, file_path: str) -> OperationResponse:
        """
        Close a document and remove from cache.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            OperationResponse with status and message
        """
        if file_path in self._documents:
            del self._documents[file_path]
            return OperationResponse.success(f"Document closed: {file_path}")
        else:
            return OperationResponse.warning(f"Document not loaded: {file_path}")
    
    def get_document_info(self, file_path: str) -> OperationResponse:
        """
        Get information about a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            OperationResponse with document information
        """
        try:
            if file_path not in self._documents:
                # Try to open the document first
                response = self.open_document(file_path, create_if_not_exists=False)
                if response.status.value != "success":
                    return response
            
            document = self._documents[file_path]
            
            # Gather document information
            tables = []
            for i, table in enumerate(document.tables):
                table_info = TableInfo(
                    index=i,
                    rows=len(table.rows),
                    columns=len(table.columns) if table.rows else 0,
                    has_headers=self._has_header_row(table),
                    style=getattr(table.style, 'name', None) if table.style else None,
                    position=i  # Simple position based on order
                )
                tables.append({
                    "index": table_info.index,
                    "rows": table_info.rows,
                    "columns": table_info.columns,
                    "has_headers": table_info.has_headers,
                    "style": table_info.style,
                    "position": table_info.position
                })
            
            data = {
                "file_path": file_path,
                "table_count": len(tables),
                "paragraph_count": len(document.paragraphs),
                "tables": tables
            }
            
            return OperationResponse.success("Document information retrieved", data)
            
        except Exception as e:
            return OperationResponse.error(f"Failed to get document info: {str(e)}")
    
    def _has_header_row(self, table) -> bool:
        """
        Check if table has a header row (simple heuristic).
        
        Args:
            table: Table object
            
        Returns:
            True if table likely has headers
        """
        if not table.rows:
            return False
        
        try:
            # Simple heuristic: if first row has different formatting or all cells have text
            first_row = table.rows[0]
            if not first_row.cells:
                return False
            
            # Check if all cells in first row have text
            has_text = all(cell.text.strip() for cell in first_row.cells)
            
            return has_text
        except:
            return False
    
    def list_loaded_documents(self) -> OperationResponse:
        """
        List all currently loaded documents.
        
        Returns:
            OperationResponse with list of loaded documents
        """
        documents = list(self._documents.keys())
        data = {
            "loaded_documents": documents,
            "count": len(documents)
        }
        
        return OperationResponse.success(f"Found {len(documents)} loaded documents", data)
