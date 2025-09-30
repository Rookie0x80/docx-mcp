"""Table operations for Word documents."""

import re
from typing import List, Optional, Dict, Any, Union
from docx import Document
from docx.table import Table, _Cell
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import qn, OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...models.responses import OperationResponse
from ...models.tables import TableInfo, CellPosition, SearchResult, TableData, TableSearchMatch, TableSearchResult
from ...models.table_analysis import (
    TableStructureAnalysis, CellStyleAnalysis, TableAnalysisResult, MergeInfo,
    CellMergeType, analyze_cell_merge, extract_cell_formatting
)
from ...models.formatting import TextFormat, CellAlignment
from ...utils.exceptions import (
    TableNotFoundError,
    InvalidTableIndexError,
    InvalidCellPositionError,
    TableOperationError,
    DataFormatError,
)
from ...utils.validation import (
    validate_table_index,
    validate_cell_position,
    validate_table_data,
    validate_position_parameter,
    sanitize_string,
)
from ...core.document_manager import DocumentManager
from .formatting import TableFormattingOperations


class TableOperations:
    """Handles table operations in Word documents."""
    
    def __init__(self, document_manager: DocumentManager):
        """
        Initialize table operations.
        
        Args:
            document_manager: Document manager instance
        """
        self.document_manager = document_manager
        self.formatting = TableFormattingOperations(document_manager)
    
    def create_table(
        self,
        file_path: str,
        rows: int,
        cols: int,
        position: str = "end",
        paragraph_index: Optional[int] = None,
        headers: Optional[List[str]] = None,
    ) -> OperationResponse:
        """
        Create a new table in the document.
        
        Args:
            file_path: Path to the document
            rows: Number of rows
            cols: Number of columns
            position: Where to insert the table
            paragraph_index: Paragraph index for 'after_paragraph' position
            headers: Optional header row data
            
        Returns:
            OperationResponse with operation result
        """
        try:
            # Validate inputs
            if rows <= 0 or cols <= 0:
                return OperationResponse.error("Rows and columns must be positive integers")
            
            valid_positions = ["end", "beginning", "after_paragraph"]
            validate_position_parameter(position, valid_positions)
            
            if position == "after_paragraph" and paragraph_index is None:
                return OperationResponse.error("paragraph_index required for 'after_paragraph' position")
            
            if headers and len(headers) != cols:
                return OperationResponse.error(f"Headers length ({len(headers)}) must match columns ({cols})")
            
            # Get document
            document = self.document_manager.get_or_load_document(file_path)
            
            # Create table
            table = None
            if position == "end":
                table = document.add_table(rows=rows, cols=cols)
            elif position == "beginning":
                # Insert at beginning by adding after title or first paragraph
                if document.paragraphs:
                    p = document.paragraphs[0]
                    table = p.insert_paragraph_before().add_table(rows=rows, cols=cols)
                else:
                    table = document.add_table(rows=rows, cols=cols)
            elif position == "after_paragraph":
                if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                    return OperationResponse.error(f"Invalid paragraph index: {paragraph_index}")
                p = document.paragraphs[paragraph_index]
                new_p = p.insert_paragraph_after()
                table = new_p._element.addnext(document.add_table(rows=rows, cols=cols)._element)
                table = document.tables[-1]  # Get the newly added table
            
            if not table:
                return OperationResponse.error("Failed to create table")
            
            # Set headers if provided
            if headers:
                for col_idx, header in enumerate(headers):
                    table.cell(0, col_idx).text = sanitize_string(header)
            
            table_index = len(document.tables) - 1
            
            data = {
                "table_index": table_index,
                "rows": rows,
                "cols": cols,
                "position": position,
                "has_headers": bool(headers)
            }
            
            return OperationResponse.success(f"Table created with {rows} rows and {cols} columns", data)
            
        except Exception as e:
            return OperationResponse.error(f"Failed to create table: {str(e)}")
    
    def delete_table(self, file_path: str, table_index: int) -> OperationResponse:
        """
        Delete a table from the document.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table to delete
            
        Returns:
            OperationResponse with operation result
        """
        try:
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            
            # Get table and remove it
            table = document.tables[table_index]
            table._element.getparent().remove(table._element)
            
            return OperationResponse.success(f"Table {table_index} deleted")
            
        except (InvalidTableIndexError, TableNotFoundError) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to delete table: {str(e)}")
    
    def add_table_rows(
        self,
        file_path: str,
        table_index: int,
        count: int = 1,
        position: str = "end",
        row_index: Optional[int] = None,
        copy_style_from_row: Optional[int] = None,
        default_text_format: Optional[TextFormat] = None,
        default_alignment: Optional[CellAlignment] = None,
        default_background_color: Optional[str] = None,
    ) -> OperationResponse:
        """
        Add rows to a table with optional styling control.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            count: Number of rows to add
            position: Where to add rows ("end", "beginning", "at_index")
            row_index: Specific row index for 'at_index' position
            copy_style_from_row: Row index to copy style from (None = no style copying)
            default_text_format: Default text formatting for new cells
            default_alignment: Default alignment for new cells
            default_background_color: Default background color for new cells
            
        Returns:
            OperationResponse with operation result
        """
        try:
            if count <= 0:
                return OperationResponse.error("Count must be a positive integer")
            
            valid_positions = ["end", "beginning", "at_index"]
            validate_position_parameter(position, valid_positions)
            
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            
            if position == "at_index" and row_index is None:
                return OperationResponse.error("row_index required for 'at_index' position")
            
            if position == "at_index":
                validate_cell_position(row_index, 0, len(table.rows), len(table.columns))
            
            # Determine reference row for style copying
            reference_row = None
            if copy_style_from_row is not None:
                # Explicit row specified
                if copy_style_from_row < 0 or copy_style_from_row >= len(table.rows):
                    return OperationResponse.error(f"Invalid copy_style_from_row: {copy_style_from_row}")
                reference_row = table.rows[copy_style_from_row]
            else:
                # Default behavior: copy from the last row (if table has rows)
                if len(table.rows) > 0:
                    if position == "end":
                        # For end insertion, copy from last row
                        reference_row = table.rows[-1]
                    elif position == "beginning":
                        # For beginning insertion, copy from first row
                        reference_row = table.rows[0]
                    elif position == "at_index" and row_index is not None:
                        # For index insertion, copy from the row at that index (or previous if available)
                        if row_index > 0:
                            reference_row = table.rows[row_index - 1]
                        elif row_index < len(table.rows):
                            reference_row = table.rows[row_index]
            
            # Keep track of newly added rows for styling
            new_rows = []
            original_row_count = len(table.rows)
            
            # Add rows
            for i in range(count):
                if position == "end":
                    new_row = table.add_row()
                    new_rows.append(new_row)
                elif position == "beginning":
                    # Insert at beginning - not directly supported, need to work around
                    new_row = table.add_row()
                    # Move new row to beginning
                    table._element.insert(i, new_row._element)
                    new_rows.append(new_row)
                elif position == "at_index":
                    # Insert at specific index
                    new_row = table.add_row()
                    # Move to desired position
                    if row_index + i < len(table.rows) - 1:
                        target_row = table.rows[row_index + i]
                        target_row._element.addprevious(new_row._element)
                    new_rows.append(new_row)
            
            # Apply styling to new rows
            self._apply_row_styling(
                new_rows, 
                reference_row, 
                default_text_format, 
                default_alignment, 
                default_background_color
            )
            
            data = {
                "table_index": table_index,
                "rows_added": count,
                "new_row_count": len(table.rows),
                "position": position
            }
            
            return OperationResponse.success(f"Added {count} rows to table {table_index}", data)
            
        except (InvalidTableIndexError, InvalidCellPositionError) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to add rows: {str(e)}")
    
    def add_table_columns(
        self,
        file_path: str,
        table_index: int,
        count: int = 1,
        position: str = "end",
        column_index: Optional[int] = None,
    ) -> OperationResponse:
        """
        Add columns to a table.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            count: Number of columns to add
            position: Where to add columns
            column_index: Specific column index for 'at_index' position
            
        Returns:
            OperationResponse with operation result
        """
        try:
            if count <= 0:
                return OperationResponse.error("Count must be a positive integer")
            
            valid_positions = ["end", "beginning", "at_index"]
            validate_position_parameter(position, valid_positions)
            
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            
            if not table.rows:
                return OperationResponse.error("Cannot add columns to empty table")
            
            original_cols = len(table.columns)
            
            if position == "at_index" and column_index is None:
                return OperationResponse.error("column_index required for 'at_index' position")
            
            if position == "at_index":
                validate_cell_position(0, column_index, len(table.rows), original_cols)
            
            # Add columns by adding cells to each row
            for _ in range(count):
                for row in table.rows:
                    if position == "end":
                        row.cells[-1]._element.addnext(row.cells[0]._element.__class__())
                    elif position == "beginning":
                        row.cells[0]._element.addprevious(row.cells[0]._element.__class__())
                    elif position == "at_index":
                        target_cell = row.cells[column_index]
                        target_cell._element.addprevious(row.cells[0]._element.__class__())
            
            new_cols = len(table.columns)
            
            data = {
                "table_index": table_index,
                "columns_added": count,
                "new_column_count": new_cols,
                "position": position
            }
            
            return OperationResponse.success(f"Added {count} columns to table {table_index}", data)
            
        except (InvalidTableIndexError, InvalidCellPositionError) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to add columns: {str(e)}")
    
    def delete_table_rows(
        self, file_path: str, table_index: int, row_indices: List[int]
    ) -> OperationResponse:
        """
        Delete rows from a table.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_indices: List of row indices to delete
            
        Returns:
            OperationResponse with operation result
        """
        try:
            if not row_indices:
                return OperationResponse.error("No row indices provided")
            
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            
            # Validate all row indices
            for row_idx in row_indices:
                validate_cell_position(row_idx, 0, len(table.rows), len(table.columns))
            
            # Sort indices in reverse order to delete from end to beginning
            sorted_indices = sorted(set(row_indices), reverse=True)
            
            # Delete rows
            for row_idx in sorted_indices:
                row = table.rows[row_idx]
                row._element.getparent().remove(row._element)
            
            data = {
                "table_index": table_index,
                "rows_deleted": len(sorted_indices),
                "remaining_rows": len(table.rows)
            }
            
            return OperationResponse.success(
                f"Deleted {len(sorted_indices)} rows from table {table_index}", data
            )
            
        except (InvalidTableIndexError, InvalidCellPositionError) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to delete rows: {str(e)}")
    
    def set_cell_value(
        self, 
        file_path: str, 
        table_index: int, 
        row_index: int, 
        column_index: int, 
        value: str,
        text_format: Optional[TextFormat] = None,
        alignment: Optional[Dict[str, str]] = None,
        background_color: Optional[str] = None,
        preserve_existing_format: bool = True
    ) -> OperationResponse:
        """
        Set the value of a specific cell with optional formatting.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_index: Row index
            column_index: Column index
            value: Value to set
            text_format: Optional text formatting (font, size, color, bold, italic, etc.)
            alignment: Optional alignment settings {"horizontal": "left/center/right", "vertical": "top/middle/bottom"}
            background_color: Optional background color as hex string (e.g., "FFFF00")
            preserve_existing_format: Whether to preserve existing formatting when not specified
            
        Returns:
            OperationResponse with operation result
        """
        try:
            from docx.shared import RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            
            validate_cell_position(row_index, column_index, len(table.rows), len(table.columns))
            
            # Get cell and set value
            cell = table.cell(row_index, column_index)
            
            # Store existing formatting if preserve_existing_format is True
            existing_format = None
            if preserve_existing_format:
                existing_format = extract_cell_formatting(cell)
            
            # Clear existing content and set new value
            cell.text = sanitize_string(value)
            
            # Apply formatting if provided
            if cell.paragraphs:
                paragraph = cell.paragraphs[0]
                
                # Apply paragraph alignment
                if alignment and alignment.get('horizontal'):
                    h_align = alignment['horizontal'].lower()
                    alignment_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if h_align in alignment_map:
                        paragraph.alignment = alignment_map[h_align]
                elif preserve_existing_format and existing_format and existing_format.get('horizontal_alignment'):
                    # Restore existing alignment
                    h_align = existing_format['horizontal_alignment']
                    alignment_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if h_align in alignment_map:
                        paragraph.alignment = alignment_map[h_align]
                
                # Apply text formatting to runs
                if paragraph.runs:
                    run = paragraph.runs[0]
                    
                    # Apply text formatting
                    if preserve_existing_format and existing_format:
                        # First restore existing text formatting
                        if existing_format.get('font_family'):
                            run.font.name = existing_format['font_family']
                        if existing_format.get('font_size'):
                            from docx.shared import Pt
                            run.font.size = Pt(existing_format['font_size'])
                        if existing_format.get('font_color'):
                            try:
                                color_hex = existing_format['font_color'].lstrip('#')
                                if len(color_hex) == 6:
                                    r = int(color_hex[0:2], 16)
                                    g = int(color_hex[2:4], 16)
                                    b = int(color_hex[4:6], 16)
                                    run.font.color.rgb = RGBColor(r, g, b)
                            except (ValueError, AttributeError):
                                pass
                        if existing_format.get('is_bold') is not None:
                            run.font.bold = existing_format['is_bold']
                        if existing_format.get('is_italic') is not None:
                            run.font.italic = existing_format['is_italic']
                        if existing_format.get('is_underlined') is not None:
                            run.font.underline = existing_format['is_underlined']
                    
                    # Then apply new text formatting (overrides existing)
                    if text_format:
                        if text_format.font_family:
                            run.font.name = text_format.font_family
                        if text_format.font_size:
                            from docx.shared import Pt
                            run.font.size = Pt(text_format.font_size)
                        if text_format.font_color:
                            # Parse hex color
                            try:
                                color_hex = text_format.font_color.lstrip('#')
                                if len(color_hex) == 6:
                                    r = int(color_hex[0:2], 16)
                                    g = int(color_hex[2:4], 16)
                                    b = int(color_hex[4:6], 16)
                                    run.font.color.rgb = RGBColor(r, g, b)
                            except (ValueError, AttributeError):
                                pass  # Skip invalid color
                        if text_format.bold is not None:
                            run.font.bold = text_format.bold
                        if text_format.italic is not None:
                            run.font.italic = text_format.italic
                        if text_format.underline is not None:
                            run.font.underline = text_format.underline
            
            # Apply vertical alignment if provided
            if alignment and alignment.get('vertical'):
                try:
                    from docx.oxml.shared import qn, OxmlElement
                    
                    v_align = alignment['vertical'].lower()
                    alignment_map = {
                        'top': 'top',
                        'middle': 'center',
                        'bottom': 'bottom'
                    }
                    
                    if v_align in alignment_map:
                        tc_pr = cell._element.get_or_add_tcPr()
                        
                        # Remove existing vAlign if present
                        existing_valign = tc_pr.find(qn('w:vAlign'))
                        if existing_valign is not None:
                            tc_pr.remove(existing_valign)
                        
                        # Add new vAlign
                        valign_element = OxmlElement('w:vAlign')
                        valign_element.set(qn('w:val'), alignment_map[v_align])
                        tc_pr.append(valign_element)
                except Exception:
                    pass  # Skip if vertical alignment application fails
            elif preserve_existing_format and existing_format and existing_format.get('vertical_alignment'):
                # Restore existing vertical alignment
                try:
                    from docx.oxml.shared import qn, OxmlElement
                    
                    v_align = existing_format['vertical_alignment'].lower()
                    alignment_map = {
                        'top': 'top',
                        'middle': 'center',
                        'bottom': 'bottom'
                    }
                    
                    if v_align in alignment_map:
                        tc_pr = cell._element.get_or_add_tcPr()
                        
                        # Remove existing vAlign if present
                        existing_valign = tc_pr.find(qn('w:vAlign'))
                        if existing_valign is not None:
                            tc_pr.remove(existing_valign)
                        
                        # Add new vAlign
                        valign_element = OxmlElement('w:vAlign')
                        valign_element.set(qn('w:val'), alignment_map[v_align])
                        tc_pr.append(valign_element)
                except Exception:
                    pass
            
            # Apply background color if provided
            if background_color:
                try:
                    # Apply cell shading using proper XML construction
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    
                    tc_pr = cell._element.get_or_add_tcPr()
                    
                    # Remove existing shading if present
                    existing_shd = tc_pr.find(qn('w:shd'))
                    if existing_shd is not None:
                        tc_pr.remove(existing_shd)
                    
                    # Create new shading element with proper namespace
                    shd_xml = f'''<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                                 w:val="clear" w:color="auto" w:fill="{background_color.lstrip('#')}"/>'''
                    shd_element = parse_xml(shd_xml)
                    tc_pr.append(shd_element)
                except Exception:
                    pass  # Skip if background color application fails
            elif preserve_existing_format and existing_format and existing_format.get('background_color'):
                # Restore existing background color
                try:
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    
                    tc_pr = cell._element.get_or_add_tcPr()
                    
                    # Remove existing shading if present
                    existing_shd = tc_pr.find(qn('w:shd'))
                    if existing_shd is not None:
                        tc_pr.remove(existing_shd)
                    
                    # Create new shading element with proper namespace
                    shd_xml = f'''<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                                 w:val="clear" w:color="auto" w:fill="{existing_format["background_color"].lstrip('#')}"/>'''
                    shd_element = parse_xml(shd_xml)
                    tc_pr.append(shd_element)
                except Exception:
                    pass
            
            # Get final formatting for response
            final_format = extract_cell_formatting(cell)
            
            data = {
                "table_index": table_index,
                "row_index": row_index,
                "column_index": column_index,
                "value": cell.text,
                "applied_formatting": {
                    "text_format": {
                        "font_family": final_format.get('font_family'),
                        "font_size": final_format.get('font_size'),
                        "font_color": final_format.get('font_color'),
                        "bold": final_format.get('is_bold', False),
                        "italic": final_format.get('is_italic', False),
                        "underlined": final_format.get('is_underlined', False)
                    },
                    "alignment": {
                        "horizontal": final_format.get('horizontal_alignment'),
                        "vertical": final_format.get('vertical_alignment')
                    },
                    "background_color": final_format.get('background_color')
                }
            }
            
            return OperationResponse.success(
                f"Cell value and formatting set at table {table_index}, row {row_index}, column {column_index}",
                data
            )
            
        except (InvalidTableIndexError, InvalidCellPositionError) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to set cell value: {str(e)}")
    
    def get_cell_value(
        self, 
        file_path: str, 
        table_index: int, 
        row_index: int, 
        column_index: int,
        include_formatting: bool = True
    ) -> OperationResponse:
        """
        Get the value and formatting of a specific cell.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_index: Row index
            column_index: Column index
            include_formatting: Whether to include detailed formatting information
            
        Returns:
            OperationResponse with cell value and formatting
        """
        try:
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            
            validate_cell_position(row_index, column_index, len(table.rows), len(table.columns))
            
            # Get cell and its value
            cell = table.cell(row_index, column_index)
            value = cell.text
            
            data = {
                "table_index": table_index,
                "row_index": row_index,
                "column_index": column_index,
                "value": value,
                "is_empty": not value.strip()
            }
            
            # Include formatting information if requested
            if include_formatting:
                cell_format = extract_cell_formatting(cell)
                merge_info = analyze_cell_merge(cell, row_index, column_index)
                
                data["formatting"] = {
                    "text_format": {
                        "font_family": cell_format.get('font_family'),
                        "font_size": cell_format.get('font_size'),
                        "font_color": cell_format.get('font_color'),
                        "bold": cell_format.get('is_bold', False),
                        "italic": cell_format.get('is_italic', False),
                        "underlined": cell_format.get('is_underlined', False),
                        "strikethrough": cell_format.get('is_strikethrough', False)
                    },
                    "alignment": {
                        "horizontal": cell_format.get('horizontal_alignment'),
                        "vertical": cell_format.get('vertical_alignment')
                    },
                    "background_color": cell_format.get('background_color'),
                    "borders": self._extract_border_data(cell_format.get('borders', {}))
                }
                
                # Include merge information if cell is merged
                if merge_info:
                    data["merge_info"] = {
                        "type": merge_info.merge_type.value,
                        "start_row": merge_info.start_row,
                        "end_row": merge_info.end_row,
                        "start_col": merge_info.start_col,
                        "end_col": merge_info.end_col,
                        "span_rows": merge_info.span_rows,
                        "span_cols": merge_info.span_cols
                    }
                else:
                    data["merge_info"] = None
            
            message = "Cell value retrieved"
            if include_formatting:
                message += " with formatting"
            
            return OperationResponse.success(message, data)
            
        except (InvalidTableIndexError, InvalidCellPositionError) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to get cell value: {str(e)}")
    
    def get_table_data(
        self,
        file_path: str,
        table_index: int,
        include_headers: bool = True,
        format_type: str = "array",
    ) -> OperationResponse:
        """
        Get all data from a table.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            include_headers: Whether to include headers
            format_type: Format of returned data ('array', 'object', 'csv')
            
        Returns:
            OperationResponse with table data
        """
        try:
            valid_formats = ["array", "object", "csv"]
            if format_type not in valid_formats:
                return OperationResponse.error(f"Invalid format. Valid options: {', '.join(valid_formats)}")
            
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            
            if not table.rows:
                return OperationResponse.success("Table is empty", {"data": []})
            
            # Extract data
            data = []
            headers = None
            
            start_row = 0
            if include_headers and table.rows:
                # Extract headers from first row
                headers = [cell.text for cell in table.rows[0].cells]
                start_row = 1
            
            # Extract data rows
            for row in table.rows[start_row:]:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
            
            # Format data according to requested format
            if format_type == "array":
                result_data = data
                if include_headers and headers:
                    result_data = [headers] + data
            elif format_type == "object":
                if headers:
                    result_data = []
                    for row in data:
                        row_dict = {}
                        for i, value in enumerate(row):
                            header = headers[i] if i < len(headers) else f"Column_{i}"
                            row_dict[header] = value
                        result_data.append(row_dict)
                else:
                    result_data = [{"Column_" + str(i): value for i, value in enumerate(row)} for row in data]
            elif format_type == "csv":
                result_data = []
                if include_headers and headers:
                    result_data.append(headers)
                result_data.extend(data)
            
            response_data = {
                "table_index": table_index,
                "format": format_type,
                "rows": len(data),
                "columns": len(data[0]) if data else 0,
                "has_headers": bool(headers),
                "headers": headers,
                "data": result_data
            }
            
            return OperationResponse.success(f"Table data retrieved in {format_type} format", response_data)
            
        except (InvalidTableIndexError,) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to get table data: {str(e)}")
    
    def list_tables(self, file_path: str, include_summary: bool = True) -> OperationResponse:
        """
        List all tables in the document.
        
        Args:
            file_path: Path to the document
            include_summary: Whether to include table summary information
            
        Returns:
            OperationResponse with list of tables
        """
        try:
            document = self.document_manager.get_or_load_document(file_path)
            
            tables = []
            for i, table in enumerate(document.tables):
                table_info = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                }
                
                if include_summary:
                    # Check if has headers (simple heuristic)
                    has_headers = False
                    if table.rows:
                        first_row_has_text = all(cell.text.strip() for cell in table.rows[0].cells)
                        has_headers = first_row_has_text
                    
                    table_info.update({
                        "has_headers": has_headers,
                        "style": getattr(table.style, 'name', None) if table.style else None,
                        "first_row_data": [cell.text for cell in table.rows[0].cells] if table.rows else []
                    })
                
                tables.append(table_info)
            
            data = {
                "tables": tables,
                "total_count": len(tables)
            }
            
            return OperationResponse.success(f"Found {len(tables)} tables", data)
            
        except Exception as e:
            return OperationResponse.error(f"Failed to list tables: {str(e)}")
    
    def search_table_content(
        self,
        file_path: str,
        query: str,
        search_mode: str = "contains",
        case_sensitive: bool = False,
        table_indices: Optional[List[int]] = None,
        max_results: Optional[int] = None
    ) -> OperationResponse:
        """
        Search for content within table cells.
        
        Args:
            file_path: Path to the document
            query: Search query string
            search_mode: Search mode ("exact", "contains", "regex")
            case_sensitive: Whether search is case sensitive
            table_indices: Optional list of table indices to search (None = all tables)
            max_results: Maximum number of results to return (None = no limit)
            
        Returns:
            OperationResponse with search results
        """
        try:
            if not query.strip():
                return OperationResponse.error("Search query cannot be empty")
            
            valid_modes = ["exact", "contains", "regex"]
            if search_mode not in valid_modes:
                return OperationResponse.error(f"Invalid search mode. Valid options: {', '.join(valid_modes)}")
            
            document = self.document_manager.get_or_load_document(file_path)
            
            # Determine which tables to search
            if table_indices is None:
                tables_to_search = list(range(len(document.tables)))
            else:
                # Validate table indices
                for idx in table_indices:
                    validate_table_index(idx, len(document.tables))
                tables_to_search = table_indices
            
            matches = []
            summary = {
                "tables_with_matches": 0,
                "matches_per_table": {},
                "total_cells_searched": 0
            }
            
            # Compile regex pattern if needed
            pattern = None
            if search_mode == "regex":
                try:
                    flags = 0 if case_sensitive else re.IGNORECASE
                    pattern = re.compile(query, flags)
                except re.error as e:
                    return OperationResponse.error(f"Invalid regex pattern: {str(e)}")
            
            # Search each table
            for table_idx in tables_to_search:
                table = document.tables[table_idx]
                table_matches = 0
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text
                        summary["total_cells_searched"] += 1
                        
                        # Perform search based on mode
                        cell_matches = self._search_cell_content(
                            cell_text, query, search_mode, case_sensitive, pattern
                        )
                        
                        # Create match objects
                        for match_info in cell_matches:
                            if max_results and len(matches) >= max_results:
                                break
                                
                            match = TableSearchMatch(
                                table_index=table_idx,
                                row_index=row_idx,
                                column_index=col_idx,
                                cell_value=cell_text,
                                match_text=match_info["text"],
                                match_start=match_info["start"],
                                match_end=match_info["end"]
                            )
                            matches.append(match)
                            table_matches += 1
                        
                        if max_results and len(matches) >= max_results:
                            break
                    
                    if max_results and len(matches) >= max_results:
                        break
                
                if table_matches > 0:
                    summary["tables_with_matches"] += 1
                    summary["matches_per_table"][table_idx] = table_matches
            
            # Create search result
            search_result = TableSearchResult(
                query=query,
                search_mode=search_mode,
                case_sensitive=case_sensitive,
                matches=matches,
                total_matches=len(matches),
                tables_searched=tables_to_search,
                summary=summary
            )
            
            message = f"Found {len(matches)} matches in {summary['tables_with_matches']} tables"
            if max_results and len(matches) >= max_results:
                message += f" (limited to {max_results} results)"
            
            return OperationResponse.success(message, search_result.to_dict())
            
        except (InvalidTableIndexError,) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to search table content: {str(e)}")
    
    def _search_cell_content(
        self,
        cell_text: str,
        query: str,
        search_mode: str,
        case_sensitive: bool,
        pattern: Optional[re.Pattern] = None
    ) -> List[Dict[str, Any]]:
        """
        Search for matches within a single cell's content.
        
        Args:
            cell_text: The cell's text content
            query: Search query
            search_mode: Search mode
            case_sensitive: Case sensitivity flag
            pattern: Compiled regex pattern (for regex mode)
            
        Returns:
            List of match information dictionaries
        """
        matches = []
        
        if not cell_text:
            return matches
        
        if search_mode == "exact":
            # Exact match
            search_text = cell_text if case_sensitive else cell_text.lower()
            query_text = query if case_sensitive else query.lower()
            
            if search_text == query_text:
                matches.append({
                    "text": cell_text,
                    "start": 0,
                    "end": len(cell_text)
                })
        
        elif search_mode == "contains":
            # Contains match
            search_text = cell_text if case_sensitive else cell_text.lower()
            query_text = query if case_sensitive else query.lower()
            
            start = 0
            while True:
                pos = search_text.find(query_text, start)
                if pos == -1:
                    break
                
                matches.append({
                    "text": cell_text[pos:pos + len(query)],
                    "start": pos,
                    "end": pos + len(query)
                })
                start = pos + 1
        
        elif search_mode == "regex":
            # Regex match
            if pattern:
                for match in pattern.finditer(cell_text):
                    matches.append({
                        "text": match.group(),
                        "start": match.start(),
                        "end": match.end()
                    })
        
        return matches
    
    def search_table_headers(
        self,
        file_path: str,
        query: str,
        search_mode: str = "contains",
        case_sensitive: bool = False
    ) -> OperationResponse:
        """
        Search specifically in table headers (first row of each table).
        
        Args:
            file_path: Path to the document
            query: Search query string
            search_mode: Search mode ("exact", "contains", "regex")
            case_sensitive: Whether search is case sensitive
            
        Returns:
            OperationResponse with search results
        """
        try:
            if not query.strip():
                return OperationResponse.error("Search query cannot be empty")
            
            document = self.document_manager.get_or_load_document(file_path)
            
            matches = []
            tables_with_headers = 0
            
            # Search only first row of each table
            for table_idx, table in enumerate(document.tables):
                if not table.rows:
                    continue
                
                first_row = table.rows[0]
                has_header_matches = False
                
                for col_idx, cell in enumerate(first_row.cells):
                    cell_text = cell.text
                    
                    # Use the same search logic as general search
                    pattern = None
                    if search_mode == "regex":
                        try:
                            flags = 0 if case_sensitive else re.IGNORECASE
                            pattern = re.compile(query, flags)
                        except re.error as e:
                            return OperationResponse.error(f"Invalid regex pattern: {str(e)}")
                    
                    cell_matches = self._search_cell_content(
                        cell_text, query, search_mode, case_sensitive, pattern
                    )
                    
                    for match_info in cell_matches:
                        match = TableSearchMatch(
                            table_index=table_idx,
                            row_index=0,  # Always first row for headers
                            column_index=col_idx,
                            cell_value=cell_text,
                            match_text=match_info["text"],
                            match_start=match_info["start"],
                            match_end=match_info["end"]
                        )
                        matches.append(match)
                        has_header_matches = True
                
                if has_header_matches:
                    tables_with_headers += 1
            
            # Create search result
            search_result = TableSearchResult(
                query=query,
                search_mode=search_mode,
                case_sensitive=case_sensitive,
                matches=matches,
                total_matches=len(matches),
                tables_searched=list(range(len(document.tables))),
                summary={
                    "search_type": "headers_only",
                    "tables_with_header_matches": tables_with_headers,
                    "total_tables": len(document.tables)
                }
            )
            
            message = f"Found {len(matches)} header matches in {tables_with_headers} tables"
            return OperationResponse.success(message, search_result.to_dict())
            
        except Exception as e:
            return OperationResponse.error(f"Failed to search table headers: {str(e)}")
    
    def analyze_table_structure(
        self,
        file_path: str,
        table_index: int,
        include_cell_details: bool = True
    ) -> OperationResponse:
        """
        Analyze the complete structure and styling of a specific table.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table to analyze
            include_cell_details: Whether to include detailed cell analysis
            
        Returns:
            OperationResponse with comprehensive table analysis
        """
        try:
            document = self.document_manager.get_or_load_document(file_path)
            
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            
            # Basic table information
            total_rows = len(table.rows)
            total_columns = len(table.columns) if table.rows else 0
            
            # Table-level properties
            table_style_name = getattr(table.style, 'name', None) if table.style else None
            
            # Header detection
            has_header_row = False
            header_row_index = None
            header_cells = None
            
            if table.rows:
                # Simple heuristic: if first row has text in all cells, consider it header
                first_row = table.rows[0]
                first_row_texts = [cell.text.strip() for cell in first_row.cells]
                has_header_row = all(text for text in first_row_texts)
                
                if has_header_row:
                    header_row_index = 0
                    header_cells = first_row_texts
            
            # Initialize cell analysis storage
            cells = []
            merge_regions = []
            merged_cells_count = 0
            
            # Style tracking for consistency analysis
            font_families = set()
            font_sizes = set()
            colors = set()
            background_colors = set()
            alignments = set()
            border_styles = set()
            
            # Analyze each cell
            for row_idx, row in enumerate(table.rows):
                cell_row = []
                for col_idx, cell in enumerate(row.cells):
                    # Extract cell content
                    text_content = cell.text
                    is_empty = not text_content.strip()
                    
                    # Analyze merge information
                    merge_info = analyze_cell_merge(cell, row_idx, col_idx)
                    if merge_info:
                        merge_regions.append(merge_info)
                        merged_cells_count += 1
                    
                    # Extract formatting if detailed analysis is requested
                    cell_analysis = None
                    if include_cell_details:
                        formatting = extract_cell_formatting(cell)
                        
                        # Track unique styles
                        if formatting["font_family"]:
                            font_families.add(formatting["font_family"])
                        if formatting["font_size"]:
                            font_sizes.add(formatting["font_size"])
                        if formatting["font_color"]:
                            colors.add(formatting["font_color"])
                        if formatting["background_color"]:
                            background_colors.add(formatting["background_color"])
                        if formatting["horizontal_alignment"]:
                            alignments.add(formatting["horizontal_alignment"])
                        
                        # Track border styles
                        for border_side, border_info in formatting["borders"].items():
                            if border_info and border_info.get("style"):
                                border_styles.add(border_info["style"])
                        
                        cell_analysis = CellStyleAnalysis(
                            row_index=row_idx,
                            column_index=col_idx,
                            text_content=text_content,
                            is_empty=is_empty,
                            merge_info=merge_info,
                            font_family=formatting["font_family"],
                            font_size=formatting["font_size"],
                            font_color=formatting["font_color"],
                            is_bold=formatting["is_bold"],
                            is_italic=formatting["is_italic"],
                            is_underlined=formatting["is_underlined"],
                            is_strikethrough=formatting["is_strikethrough"],
                            horizontal_alignment=formatting["horizontal_alignment"],
                            vertical_alignment=formatting["vertical_alignment"],
                            background_color=formatting["background_color"],
                            top_border=formatting["borders"]["top"],
                            bottom_border=formatting["borders"]["bottom"],
                            left_border=formatting["borders"]["left"],
                            right_border=formatting["borders"]["right"],
                            width=None,  # Could be implemented if needed
                            height=None  # Could be implemented if needed
                        )
                    else:
                        # Minimal cell analysis without formatting details
                        cell_analysis = CellStyleAnalysis(
                            row_index=row_idx,
                            column_index=col_idx,
                            text_content=text_content,
                            is_empty=is_empty,
                            merge_info=merge_info,
                            font_family=None,
                            font_size=None,
                            font_color=None,
                            is_bold=False,
                            is_italic=False,
                            is_underlined=False,
                            is_strikethrough=False,
                            horizontal_alignment=None,
                            vertical_alignment=None,
                            background_color=None,
                            top_border=None,
                            bottom_border=None,
                            left_border=None,
                            right_border=None,
                            width=None,
                            height=None
                        )
                    
                    cell_row.append(cell_analysis)
                
                cells.append(cell_row)
            
            # Style consistency analysis
            consistent_fonts = len(font_families) <= 1
            consistent_alignment = len(alignments) <= 1
            consistent_borders = len(border_styles) <= 1
            
            # Create table structure analysis
            table_analysis = TableStructureAnalysis(
                table_index=table_index,
                total_rows=total_rows,
                total_columns=total_columns,
                table_style_name=table_style_name,
                table_alignment=None,  # Could be implemented if needed
                table_width=None,      # Could be implemented if needed
                has_header_row=has_header_row,
                header_row_index=header_row_index,
                header_cells=header_cells,
                cells=cells,
                merged_cells_count=merged_cells_count,
                merge_regions=merge_regions,
                consistent_fonts=consistent_fonts,
                consistent_alignment=consistent_alignment,
                consistent_borders=consistent_borders,
                unique_font_families=list(font_families),
                unique_font_sizes=list(font_sizes),
                unique_colors=list(colors),
                unique_background_colors=list(background_colors)
            )
            
            return OperationResponse.success(
                f"Table {table_index} structure analyzed successfully",
                table_analysis.to_dict()
            )
            
        except (InvalidTableIndexError,) as e:
            return OperationResponse.error(str(e))
        except Exception as e:
            return OperationResponse.error(f"Failed to analyze table structure: {str(e)}")
    
    def analyze_all_tables(
        self,
        file_path: str,
        include_cell_details: bool = True
    ) -> OperationResponse:
        """
        Analyze the structure and styling of all tables in the document.
        
        Args:
            file_path: Path to the document
            include_cell_details: Whether to include detailed cell analysis
            
        Returns:
            OperationResponse with analysis of all tables
        """
        try:
            from datetime import datetime
            
            document = self.document_manager.get_or_load_document(file_path)
            
            if not document.tables:
                return OperationResponse.success(
                    "No tables found in document",
                    {"file_path": file_path, "total_tables": 0, "tables": []}
                )
            
            table_analyses = []
            
            # Analyze each table
            for table_idx in range(len(document.tables)):
                analysis_response = self.analyze_table_structure(
                    file_path, table_idx, include_cell_details
                )
                
                if analysis_response.success:
                    # Extract the table analysis from the response data
                    table_data = analysis_response.data
                    table_analyses.append(TableStructureAnalysis(
                        table_index=table_data["table_info"]["index"],
                        total_rows=table_data["table_info"]["rows"],
                        total_columns=table_data["table_info"]["columns"],
                        table_style_name=table_data["table_info"]["style_name"],
                        table_alignment=table_data["table_info"]["alignment"],
                        table_width=table_data["table_info"]["width"],
                        has_header_row=table_data["header_info"]["has_header"],
                        header_row_index=table_data["header_info"]["header_row_index"],
                        header_cells=table_data["header_info"]["header_cells"],
                        cells=[],  # We'll populate this if needed
                        merged_cells_count=table_data["merge_analysis"]["merged_cells_count"],
                        merge_regions=[],  # We'll populate this if needed
                        consistent_fonts=table_data["style_consistency"]["fonts"],
                        consistent_alignment=table_data["style_consistency"]["alignment"],
                        consistent_borders=table_data["style_consistency"]["borders"],
                        unique_font_families=table_data["style_summary"]["font_families"],
                        unique_font_sizes=table_data["style_summary"]["font_sizes"],
                        unique_colors=table_data["style_summary"]["colors"],
                        unique_background_colors=table_data["style_summary"]["background_colors"]
                    ))
                else:
                    # If individual table analysis fails, log it but continue
                    continue
            
            # Create comprehensive analysis result
            analysis_result = TableAnalysisResult(
                file_path=file_path,
                total_tables=len(table_analyses),
                analysis_timestamp=datetime.now().isoformat(),
                tables=table_analyses
            )
            
            return OperationResponse.success(
                f"Analyzed {len(table_analyses)} tables successfully",
                analysis_result.to_dict()
            )
            
        except Exception as e:
            return OperationResponse.error(f"Failed to analyze all tables: {str(e)}")
    
    def _apply_row_styling(
        self, 
        new_rows, 
        reference_row, 
        default_text_format, 
        default_alignment, 
        default_background_color
    ):
        """
        Apply styling to newly added rows.
        
        Args:
            new_rows: List of newly created row objects
            reference_row: Row to copy style from (if provided)
            default_text_format: Default text formatting
            default_alignment: Default alignment
            default_background_color: Default background color
        """
        
        for new_row in new_rows:
            # Apply styling to each cell in the new row
            for col_idx, new_cell in enumerate(new_row.cells):
                # Determine reference cell for style copying
                reference_cell = None
                if reference_row and col_idx < len(reference_row.cells):
                    reference_cell = reference_row.cells[col_idx]
                
                # Copy style from reference cell if available
                if reference_cell:
                    self._copy_cell_style(new_cell, reference_cell)
                
                # Apply default formatting if no reference or to override
                if default_text_format or default_alignment or default_background_color:
                    self._apply_default_cell_formatting(
                        new_cell, 
                        default_text_format, 
                        default_alignment, 
                        default_background_color
                    )
    
    def _copy_cell_style(self, target_cell, source_cell):
        """
        Copy all styling from source cell to target cell.
        
        Args:
            target_cell: Cell to apply styling to
            source_cell: Cell to copy styling from
        """
        try:
            # Copy paragraph formatting
            for target_para, source_para in zip(target_cell.paragraphs, source_cell.paragraphs):
                # Copy paragraph alignment
                target_para.alignment = source_para.alignment
                
                # Copy run formatting
                if source_para.runs:
                    # Clear existing runs in target
                    for run in target_para.runs:
                        run._element.getparent().remove(run._element)
                    
                    # Copy runs from source
                    for source_run in source_para.runs:
                        new_run = target_para.add_run("")
                        
                        # Copy font properties
                        if source_run.font.name:
                            new_run.font.name = source_run.font.name
                        if source_run.font.size:
                            new_run.font.size = source_run.font.size
                        if source_run.font.color.rgb:
                            new_run.font.color.rgb = source_run.font.color.rgb
                        
                        new_run.bold = source_run.bold
                        new_run.italic = source_run.italic
                        new_run.underline = source_run.underline
            
            # Copy cell background color
            try:
                source_shading = source_cell._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if source_shading is not None:
                    target_shading = target_cell._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    if target_shading is None:
                        # Create shading element
                        target_shading = OxmlElement('w:shd')
                        target_cell._element.get_or_add_tcPr().append(target_shading)
                    
                    # Copy fill attribute
                    if source_shading.get(qn('w:fill')):
                        target_shading.set(qn('w:fill'), source_shading.get(qn('w:fill')))
            except Exception:
                pass  # Ignore background color copy errors
            
            # Copy cell vertical alignment
            try:
                source_valign = source_cell._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign')
                if source_valign is not None:
                    target_valign = target_cell._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign')
                    if target_valign is None:
                        target_valign = OxmlElement('w:vAlign')
                        target_cell._element.get_or_add_tcPr().append(target_valign)
                    target_valign.set(qn('w:val'), source_valign.get(qn('w:val')))
            except Exception:
                pass  # Ignore vertical alignment copy errors
            
            # Copy cell borders
            try:
                ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                source_borders = source_cell._element.find(f'.//{{{ns}}}tcBorders')
                if source_borders is not None:
                    # Get or create target tcPr
                    target_tcPr = target_cell._element.get_or_add_tcPr()
                    
                    # Remove existing borders
                    existing_borders = target_tcPr.find(f'.//{{{ns}}}tcBorders')
                    if existing_borders is not None:
                        target_tcPr.remove(existing_borders)
                    
                    # Clone the entire tcBorders element
                    import copy
                    new_borders = copy.deepcopy(source_borders)
                    target_tcPr.append(new_borders)
            except Exception:
                pass  # Ignore border copy errors
                
        except Exception as e:
            # If copying fails, just continue - better to have unstyled cells than no cells
            pass
    
    def _apply_default_cell_formatting(
        self, 
        cell, 
        text_format, 
        alignment, 
        background_color
    ):
        """
        Apply default formatting to a cell.
        
        Args:
            cell: Cell to format
            text_format: TextFormat object with formatting
            alignment: CellAlignment object with alignment
            background_color: Background color as hex string
        """
        try:
            # Apply text formatting
            if text_format:
                for paragraph in cell.paragraphs:
                    if not paragraph.runs:
                        paragraph.add_run("")
                    
                    for run in paragraph.runs:
                        if text_format.font_family:
                            run.font.name = text_format.font_family
                        if text_format.font_size:
                            run.font.size = Pt(text_format.font_size)
                        if text_format.font_color:
                            # Parse hex color
                            try:
                                color_hex = text_format.font_color.lstrip('#')
                                if len(color_hex) == 6:
                                    r = int(color_hex[0:2], 16)
                                    g = int(color_hex[2:4], 16)
                                    b = int(color_hex[4:6], 16)
                                    run.font.color.rgb = RGBColor(r, g, b)
                            except Exception:
                                pass
                        
                        if text_format.bold is not None:
                            run.bold = text_format.bold
                        if text_format.italic is not None:
                            run.italic = text_format.italic
                        if text_format.underline is not None:
                            run.underline = text_format.underline
            
            # Apply alignment
            if alignment:
                for paragraph in cell.paragraphs:
                    if alignment.horizontal:
                        alignment_map = {
                            'left': WD_ALIGN_PARAGRAPH.LEFT,
                            'center': WD_ALIGN_PARAGRAPH.CENTER,
                            'right': WD_ALIGN_PARAGRAPH.RIGHT,
                            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                        }
                        if alignment.horizontal.lower() in alignment_map:
                            paragraph.alignment = alignment_map[alignment.horizontal.lower()]
                
                # Apply vertical alignment
                if alignment.vertical:
                    try:
                        v_align = alignment.vertical.lower()
                        if v_align == 'middle':
                            v_align = 'center'
                        
                        valign_element = cell._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign')
                        if valign_element is None:
                            valign_element = OxmlElement('w:vAlign')
                            cell._element.get_or_add_tcPr().append(valign_element)
                        valign_element.set(qn('w:val'), v_align)
                    except Exception:
                        pass
            
            # Apply background color
            if background_color:
                try:
                    color_hex = background_color.lstrip('#').upper()
                    if len(color_hex) == 6:
                        shading = cell._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                        if shading is None:
                            shading = OxmlElement('w:shd')
                            cell._element.get_or_add_tcPr().append(shading)
                        shading.set(qn('w:fill'), color_hex)
                except Exception:
                    pass
                    
        except Exception as e:
            # If formatting fails, continue - better to have unformatted cells than no cells
            pass
    
    def _extract_border_data(self, borders_dict: dict) -> dict:
        """
        Extract border data from extract_cell_formatting format to expected format.
        
        Args:
            borders_dict: Border data from extract_cell_formatting
            
        Returns:
            Border data in expected format for get_cell_value
        """
        result = {}
        
        for side in ['top', 'bottom', 'left', 'right']:
            border_info = borders_dict.get(side, {})
            if border_info:
                result[f'{side}_style'] = border_info.get('style')
                result[f'{side}_width'] = border_info.get('width')
                result[f'{side}_color'] = border_info.get('color')
            else:
                result[f'{side}_style'] = None
                result[f'{side}_width'] = None
                result[f'{side}_color'] = None
        
        return result
