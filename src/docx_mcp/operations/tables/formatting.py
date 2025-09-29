"""Table and cell formatting operations."""

from typing import Optional, Dict, Any, Union
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from ...models.responses import OperationResponse
from ...models.formatting import (
    TextFormat, CellAlignment, CellBorders, CellFormatting,
    HorizontalAlignment, VerticalAlignment, BorderStyle, BorderWidth,
    hex_to_rgb, validate_color
)
from ...utils.exceptions import (
    TableNotFoundError, InvalidCellPositionError, 
    TableOperationError, DataFormatError
)
from ...utils.validation import validate_table_index, validate_cell_position


class TableFormattingOperations:
    """Handles table and cell formatting operations."""
    
    def __init__(self, document_manager):
        """Initialize with document manager."""
        self.document_manager = document_manager

    def format_cell_text(
        self,
        file_path: str,
        table_index: int,
        row_index: int,
        column_index: int,
        text_format: Union[TextFormat, Dict[str, Any]]
    ) -> OperationResponse:
        """
        Apply text formatting to a specific cell.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_index: Row index of the cell
            column_index: Column index of the cell
            text_format: TextFormat object or dictionary with formatting options
            
        Returns:
            OperationResponse indicating success or failure
        """
        try:
            document = self.document_manager.get_document(file_path)
            if not document:
                return OperationResponse.error(f"Document not loaded: {file_path}")
            
            # Validate parameters
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            validate_cell_position(row_index, column_index, len(table.rows), len(table.columns))
            
            # Convert dict to TextFormat if needed
            if isinstance(text_format, dict):
                text_format = TextFormat.from_dict(text_format)
            
            # Get the cell
            cell = table.rows[row_index].cells[column_index]
            
            # Apply formatting to all paragraphs in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._apply_text_formatting(run, text_format)
                
                # If no runs exist, create one
                if not paragraph.runs:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    self._apply_text_formatting(run, text_format)
            
            return OperationResponse.success(
                f"Text formatting applied to cell [{row_index}, {column_index}]",
                {
                    "table_index": table_index,
                    "row_index": row_index,
                    "column_index": column_index,
                    "formatting_applied": text_format.to_dict()
                }
            )
            
        except Exception as e:
            return OperationResponse.error(f"Failed to format cell text: {str(e)}")

    def format_cell_alignment(
        self,
        file_path: str,
        table_index: int,
        row_index: int,
        column_index: int,
        alignment: Union[CellAlignment, Dict[str, Any]]
    ) -> OperationResponse:
        """
        Set text alignment for a specific cell.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_index: Row index of the cell
            column_index: Column index of the cell
            alignment: CellAlignment object or dictionary with alignment settings
            
        Returns:
            OperationResponse indicating success or failure
        """
        try:
            document = self.document_manager.get_document(file_path)
            if not document:
                return OperationResponse.error(f"Document not loaded: {file_path}")
            
            # Validate parameters
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            validate_cell_position(row_index, column_index, len(table.rows), len(table.columns))
            
            # Convert dict to CellAlignment if needed
            if isinstance(alignment, dict):
                alignment = CellAlignment.from_dict(alignment)
            
            # Get the cell
            cell = table.rows[row_index].cells[column_index]
            
            # Apply horizontal alignment to paragraphs
            if alignment.horizontal:
                wd_alignment = self._get_paragraph_alignment(alignment.horizontal)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = wd_alignment
            
            # Apply vertical alignment to cell
            if alignment.vertical:
                self._set_cell_vertical_alignment(cell, alignment.vertical)
            
            return OperationResponse.success(
                f"Alignment applied to cell [{row_index}, {column_index}]",
                {
                    "table_index": table_index,
                    "row_index": row_index,
                    "column_index": column_index,
                    "alignment_applied": alignment.to_dict()
                }
            )
            
        except Exception as e:
            return OperationResponse.error(f"Failed to set cell alignment: {str(e)}")

    def format_cell_background(
        self,
        file_path: str,
        table_index: int,
        row_index: int,
        column_index: int,
        color: str
    ) -> OperationResponse:
        """
        Set background color for a specific cell.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_index: Row index of the cell
            column_index: Column index of the cell
            color: Hex color string (with or without #)
            
        Returns:
            OperationResponse indicating success or failure
        """
        try:
            document = self.document_manager.get_document(file_path)
            if not document:
                return OperationResponse.error(f"Document not loaded: {file_path}")
            
            # Validate parameters
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            validate_cell_position(row_index, column_index, len(table.rows), len(table.columns))
            
            # Validate and clean color
            color = color.lstrip('#')
            if not validate_color(color):
                return OperationResponse.error(f"Invalid color format: {color}")
            
            # Get the cell
            cell = table.rows[row_index].cells[column_index]
            
            # Set background color using shading
            self._set_cell_background_color(cell, color)
            
            return OperationResponse.success(
                f"Background color applied to cell [{row_index}, {column_index}]",
                {
                    "table_index": table_index,
                    "row_index": row_index,
                    "column_index": column_index,
                    "background_color": color
                }
            )
            
        except Exception as e:
            return OperationResponse.error(f"Failed to set cell background: {str(e)}")

    def format_cell_borders(
        self,
        file_path: str,
        table_index: int,
        row_index: int,
        column_index: int,
        borders: Union[CellBorders, Dict[str, Any]]
    ) -> OperationResponse:
        """
        Set borders for a specific cell.
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_index: Row index of the cell
            column_index: Column index of the cell
            borders: CellBorders object or dictionary with border settings
            
        Returns:
            OperationResponse indicating success or failure
        """
        try:
            document = self.document_manager.get_document(file_path)
            if not document:
                return OperationResponse.error(f"Document not loaded: {file_path}")
            
            # Validate parameters
            validate_table_index(table_index, len(document.tables))
            table = document.tables[table_index]
            validate_cell_position(row_index, column_index, len(table.rows), len(table.columns))
            
            # Convert dict to CellBorders if needed
            if isinstance(borders, dict):
                borders = CellBorders.from_dict(borders)
            
            # Get the cell
            cell = table.rows[row_index].cells[column_index]
            
            # Apply borders
            self._set_cell_borders(cell, borders)
            
            return OperationResponse.success(
                f"Borders applied to cell [{row_index}, {column_index}]",
                {
                    "table_index": table_index,
                    "row_index": row_index,
                    "column_index": column_index,
                    "borders_applied": borders.to_dict()
                }
            )
            
        except Exception as e:
            return OperationResponse.error(f"Failed to set cell borders: {str(e)}")

    def format_cell_complete(
        self,
        file_path: str,
        table_index: int,
        row_index: int,
        column_index: int,
        formatting: Union[CellFormatting, Dict[str, Any]]
    ) -> OperationResponse:
        """
        Apply complete formatting to a cell (text, alignment, background, borders).
        
        Args:
            file_path: Path to the document
            table_index: Index of the table
            row_index: Row index of the cell
            column_index: Column index of the cell
            formatting: CellFormatting object or dictionary with all formatting options
            
        Returns:
            OperationResponse indicating success or failure
        """
        try:
            # Convert dict to CellFormatting if needed
            if isinstance(formatting, dict):
                formatting = CellFormatting.from_dict(formatting)
            
            results = []
            
            # Apply text formatting
            if formatting.text_format:
                result = self.format_cell_text(
                    file_path, table_index, row_index, column_index, formatting.text_format
                )
                if result.status.value != "success":
                    return result
                results.append("text_format")
            
            # Apply alignment
            if formatting.alignment:
                result = self.format_cell_alignment(
                    file_path, table_index, row_index, column_index, formatting.alignment
                )
                if result.status.value != "success":
                    return result
                results.append("alignment")
            
            # Apply background color
            if formatting.background_color:
                result = self.format_cell_background(
                    file_path, table_index, row_index, column_index, formatting.background_color
                )
                if result.status.value != "success":
                    return result
                results.append("background_color")
            
            # Apply borders
            if formatting.borders:
                result = self.format_cell_borders(
                    file_path, table_index, row_index, column_index, formatting.borders
                )
                if result.status.value != "success":
                    return result
                results.append("borders")
            
            return OperationResponse.success(
                f"Complete formatting applied to cell [{row_index}, {column_index}]",
                {
                    "table_index": table_index,
                    "row_index": row_index,
                    "column_index": column_index,
                    "applied_formats": results,
                    "formatting": formatting.to_dict()
                }
            )
            
        except Exception as e:
            return OperationResponse.error(f"Failed to apply complete cell formatting: {str(e)}")

    # Helper methods
    
    def _apply_text_formatting(self, run, text_format: TextFormat):
        """Apply text formatting to a run."""
        if text_format.font_family:
            run.font.name = text_format.font_family
        
        if text_format.font_size:
            run.font.size = Pt(text_format.font_size)
        
        if text_format.font_color:
            color = text_format.font_color.lstrip('#')
            if validate_color(color):
                r, g, b = hex_to_rgb(color)
                run.font.color.rgb = RGBColor(r, g, b)
        
        if text_format.bold is not None:
            run.font.bold = text_format.bold
        
        if text_format.italic is not None:
            run.font.italic = text_format.italic
        
        if text_format.underline is not None:
            run.font.underline = text_format.underline
        
        if text_format.strikethrough is not None:
            run.font.strike = text_format.strikethrough
        
        if text_format.subscript is not None:
            run.font.subscript = text_format.subscript
        
        if text_format.superscript is not None:
            run.font.superscript = text_format.superscript

    def _get_paragraph_alignment(self, alignment: HorizontalAlignment):
        """Convert HorizontalAlignment to WD_PARAGRAPH_ALIGNMENT."""
        alignment_map = {
            HorizontalAlignment.LEFT: WD_PARAGRAPH_ALIGNMENT.LEFT,
            HorizontalAlignment.CENTER: WD_PARAGRAPH_ALIGNMENT.CENTER,
            HorizontalAlignment.RIGHT: WD_PARAGRAPH_ALIGNMENT.RIGHT,
            HorizontalAlignment.JUSTIFY: WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        return alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def _set_cell_vertical_alignment(self, cell, alignment: VerticalAlignment):
        """Set vertical alignment for a cell."""
        # This requires direct XML manipulation as python-docx doesn't have direct support
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove existing vAlign if present
        for vAlign in tcPr.xpath('.//w:vAlign'):
            vAlign.getparent().remove(vAlign)
        
        # Add new vAlign
        vAlign = OxmlElement('w:vAlign')
        alignment_map = {
            VerticalAlignment.TOP: "top",
            VerticalAlignment.MIDDLE: "center", 
            VerticalAlignment.BOTTOM: "bottom"
        }
        vAlign.set(qn('w:val'), alignment_map.get(alignment, "top"))
        tcPr.append(vAlign)

    def _set_cell_background_color(self, cell, color: str):
        """Set background color for a cell."""
        # Direct XML manipulation for cell shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove existing shd if present
        for shd in tcPr.xpath('.//w:shd'):
            shd.getparent().remove(shd)
        
        # Add new shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color.upper())
        tcPr.append(shd)

    def _set_cell_borders(self, cell, borders: CellBorders):
        """Set borders for a cell."""
        # Direct XML manipulation for cell borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Remove existing tcBorders if present
        for tcBorders in tcPr.xpath('.//w:tcBorders'):
            tcBorders.getparent().remove(tcBorders)
        
        # Add new borders
        tcBorders = OxmlElement('w:tcBorders')
        
        border_sides = {
            'top': borders.top,
            'bottom': borders.bottom,
            'left': borders.left,
            'right': borders.right
        }
        
        for side_name, border_props in border_sides.items():
            if border_props:
                border_element = OxmlElement(f'w:{side_name}')
                
                # Set border style
                style_map = {
                    BorderStyle.NONE: "none",
                    BorderStyle.SOLID: "single",
                    BorderStyle.DASHED: "dashed",
                    BorderStyle.DOTTED: "dotted",
                    BorderStyle.DOUBLE: "double"
                }
                border_element.set(qn('w:val'), style_map.get(border_props.style, "single"))
                
                # Set border width
                width_map = {
                    BorderWidth.THIN: "4",
                    BorderWidth.MEDIUM: "8", 
                    BorderWidth.THICK: "12"
                }
                border_element.set(qn('w:sz'), width_map.get(border_props.width, "4"))
                
                # Set border color
                border_element.set(qn('w:color'), border_props.color.upper())
                
                tcBorders.append(border_element)
        
        tcPr.append(tcBorders)
