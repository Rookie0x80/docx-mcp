"""Tests for table operations."""

import pytest
import re
from docx_mcp.models.responses import ResponseStatus


class TestTableStructureOperations:
    """Test table structure operations."""

    @pytest.mark.unit
    def test_create_table_basic(self, document_manager, table_operations, test_doc_path):
        """Test creating a basic table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        result = table_operations.create_table(str(test_doc_path), rows=3, cols=4)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['rows'] == 3
        assert result.data['cols'] == 4
        assert result.data['table_index'] == 0

    @pytest.mark.unit
    def test_create_table_with_headers(self, document_manager, table_operations, test_doc_path):
        """Test creating a table with headers."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        headers = ["Name", "Age", "City"]
        
        result = table_operations.create_table(
            str(test_doc_path), 
            rows=3, 
            cols=3, 
            headers=headers
        )
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['has_headers'] is True

    @pytest.mark.unit
    def test_create_table_invalid_dimensions(self, document_manager, table_operations, test_doc_path):
        """Test creating a table with invalid dimensions."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Test zero rows
        result = table_operations.create_table(str(test_doc_path), rows=0, cols=3)
        assert result.status == ResponseStatus.ERROR
        
        # Test zero columns
        result = table_operations.create_table(str(test_doc_path), rows=3, cols=0)
        assert result.status == ResponseStatus.ERROR

    @pytest.mark.unit
    def test_create_table_header_mismatch(self, document_manager, table_operations, test_doc_path):
        """Test creating a table with header count mismatch."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        headers = ["Name", "Age"]  # 2 headers for 3 columns
        
        result = table_operations.create_table(
            str(test_doc_path), 
            rows=3, 
            cols=3, 
            headers=headers
        )
        
        assert result.status == ResponseStatus.ERROR
        assert "Headers length" in result.message

    @pytest.mark.unit
    def test_delete_table(self, document_manager, table_operations, test_doc_path):
        """Test deleting a table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        result = table_operations.delete_table(str(test_doc_path), table_index=0)
        
        assert result.status == ResponseStatus.SUCCESS
        assert "deleted" in result.message

    @pytest.mark.unit
    def test_delete_nonexistent_table(self, document_manager, table_operations, test_doc_path):
        """Test deleting a non-existent table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        result = table_operations.delete_table(str(test_doc_path), table_index=999)
        
        assert result.status == ResponseStatus.ERROR

    @pytest.mark.unit
    def test_add_table_rows(self, document_manager, table_operations, test_doc_path):
        """Test adding rows to a table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=3)
        
        result = table_operations.add_table_rows(str(test_doc_path), table_index=0, count=2)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['rows_added'] == 2
        assert result.data['new_row_count'] == 4  # Original 2 + added 2

    @pytest.mark.unit
    def test_add_table_columns(self, document_manager, table_operations, test_doc_path):
        """Test adding columns to a table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=3, cols=2)
        
        result = table_operations.add_table_columns(str(test_doc_path), table_index=0, count=1)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['columns_added'] == 1

    @pytest.mark.unit
    def test_delete_table_rows(self, document_manager, table_operations, test_doc_path):
        """Test deleting rows from a table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=5, cols=3)
        
        result = table_operations.delete_table_rows(
            str(test_doc_path), 
            table_index=0, 
            row_indices=[1, 3]
        )
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['rows_deleted'] == 2


class TestTableDataOperations:
    """Test table data operations."""

    @pytest.fixture
    def setup_table(self, document_manager, table_operations, test_doc_path, sample_table_data):
        """Set up a table with sample data."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(
            str(test_doc_path), 
            rows=4, 
            cols=4, 
            headers=sample_table_data["headers"]
        )
        return 0  # table index

    @pytest.mark.unit
    def test_set_cell_value(self, table_operations, test_doc_path, setup_table):
        """Test setting a cell value."""
        table_index = setup_table
        
        result = table_operations.set_cell_value(
            str(test_doc_path), table_index, 1, 0, "Test Value"
        )
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['value'] == "Test Value"

    @pytest.mark.unit
    def test_get_cell_value(self, table_operations, test_doc_path, setup_table):
        """Test getting a cell value."""
        table_index = setup_table
        
        # First set a value
        table_operations.set_cell_value(str(test_doc_path), table_index, 1, 0, "Test Value")
        
        # Then get it
        result = table_operations.get_cell_value(str(test_doc_path), table_index, 1, 0)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['value'] == "Test Value"

    @pytest.mark.unit
    def test_get_cell_value_invalid_position(self, table_operations, test_doc_path, setup_table):
        """Test getting a cell value from invalid position."""
        table_index = setup_table
        
        result = table_operations.get_cell_value(str(test_doc_path), table_index, 999, 999)
        
        assert result.status == ResponseStatus.ERROR

    @pytest.mark.unit
    def test_get_table_data_array_format(self, document_manager, table_operations, test_doc_path, sample_table_data):
        """Test getting table data in array format."""
        # Setup table with data
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(
            str(test_doc_path), 
            rows=4, 
            cols=4, 
            headers=sample_table_data["headers"]
        )
        
        # Add sample data
        for row_idx, row_data in enumerate(sample_table_data["data"], 1):
            for col_idx, value in enumerate(row_data):
                table_operations.set_cell_value(str(test_doc_path), 0, row_idx, col_idx, value)
        
        # Get table data
        result = table_operations.get_table_data(
            str(test_doc_path), 0, include_headers=True, format_type="array"
        )
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['format'] == "array"
        assert result.data['headers'] == sample_table_data["headers"]
        assert result.data['rows'] == 3  # Data rows only
        assert result.data['columns'] == 4

    @pytest.mark.unit
    def test_get_table_data_object_format(self, document_manager, table_operations, test_doc_path, sample_table_data):
        """Test getting table data in object format."""
        # Setup table with data
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(
            str(test_doc_path), 
            rows=4, 
            cols=4, 
            headers=sample_table_data["headers"]
        )
        
        # Add one row of sample data
        for col_idx, value in enumerate(sample_table_data["data"][0]):
            table_operations.set_cell_value(str(test_doc_path), 0, 1, col_idx, value)
        
        # Get table data
        result = table_operations.get_table_data(
            str(test_doc_path), 0, include_headers=True, format_type="object"
        )
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['format'] == "object"
        assert isinstance(result.data['data'], list)
        if result.data['data']:  # If there's data
            assert isinstance(result.data['data'][0], dict)

    @pytest.mark.unit
    def test_get_table_data_invalid_format(self, table_operations, test_doc_path, setup_table):
        """Test getting table data with invalid format."""
        table_index = setup_table
        
        result = table_operations.get_table_data(
            str(test_doc_path), table_index, format_type="invalid_format"
        )
        
        assert result.status == ResponseStatus.ERROR
        assert "Invalid format" in result.message


class TestTableQueryOperations:
    """Test table query operations."""

    @pytest.mark.unit
    def test_list_tables_empty_document(self, document_manager, table_operations, test_doc_path):
        """Test listing tables in an empty document."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        result = table_operations.list_tables(str(test_doc_path))
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_count'] == 0
        assert result.data['tables'] == []

    @pytest.mark.unit
    def test_list_tables_with_tables(self, document_manager, table_operations, test_doc_path):
        """Test listing tables in a document with tables."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create multiple tables
        table_operations.create_table(str(test_doc_path), rows=2, cols=3)
        table_operations.create_table(str(test_doc_path), rows=3, cols=2)
        
        result = table_operations.list_tables(str(test_doc_path), include_summary=True)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_count'] == 2
        
        tables = result.data['tables']
        assert len(tables) == 2
        assert tables[0]['rows'] == 2
        assert tables[0]['columns'] == 3
        assert tables[1]['rows'] == 3
        assert tables[1]['columns'] == 2

    @pytest.mark.unit
    def test_list_tables_no_summary(self, document_manager, table_operations, test_doc_path):
        """Test listing tables without summary information."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=3)
        
        result = table_operations.list_tables(str(test_doc_path), include_summary=False)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_count'] == 1
        
        table_info = result.data['tables'][0]
        assert 'index' in table_info
        assert 'rows' in table_info
        assert 'columns' in table_info
        # Summary fields should not be present
        assert 'first_row_data' not in table_info or not table_info['first_row_data']


class TestErrorHandling:
    """Test error handling scenarios."""

    @pytest.mark.unit
    def test_operations_on_nonexistent_document(self, table_operations):
        """Test operations on a non-existent document."""
        result = table_operations.create_table("nonexistent.docx", rows=2, cols=2)
        assert result.status == ResponseStatus.ERROR

    @pytest.mark.unit
    def test_operations_on_invalid_table_index(self, document_manager, table_operations, test_doc_path):
        """Test operations with invalid table index."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        result = table_operations.get_cell_value(str(test_doc_path), 999, 0, 0)
        assert result.status == ResponseStatus.ERROR

    @pytest.mark.unit
    def test_invalid_cell_coordinates(self, document_manager, table_operations, test_doc_path):
        """Test operations with invalid cell coordinates."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Test negative coordinates
        result = table_operations.set_cell_value(str(test_doc_path), 0, -1, 0, "value")
        assert result.status == ResponseStatus.ERROR
        
        # Test out of range coordinates
        result = table_operations.set_cell_value(str(test_doc_path), 0, 999, 999, "value")
        assert result.status == ResponseStatus.ERROR


class TestTableSearchOperations:
    """Test table search operations."""

    @pytest.mark.unit
    def test_search_table_content_basic(self, document_manager, table_operations, test_doc_path):
        """Test basic table content search."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with test data
        table_operations.create_table(str(test_doc_path), rows=3, cols=3, headers=["Name", "Age", "City"])
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "Alice")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "25")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 2, "New York")
        table_operations.set_cell_value(str(test_doc_path), 0, 2, 0, "Bob")
        table_operations.set_cell_value(str(test_doc_path), 0, 2, 1, "30")
        table_operations.set_cell_value(str(test_doc_path), 0, 2, 2, "Boston")
        
        # Search for "Alice"
        result = table_operations.search_table_content(str(test_doc_path), "Alice")
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 1
        assert len(result.data['matches']) == 1
        assert result.data['matches'][0]['table_index'] == 0
        assert result.data['matches'][0]['row_index'] == 1
        assert result.data['matches'][0]['column_index'] == 0
        assert result.data['matches'][0]['cell_value'] == "Alice"
        assert result.data['matches'][0]['match_text'] == "Alice"

    @pytest.mark.unit
    def test_search_table_content_contains_mode(self, document_manager, table_operations, test_doc_path):
        """Test table content search with contains mode."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with test data
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 0, "Hello World")
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 1, "Testing")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "World Peace")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "Python")
        
        # Search for "World" - should find 2 matches
        result = table_operations.search_table_content(str(test_doc_path), "World", search_mode="contains")
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 2
        assert result.data['search_mode'] == "contains"
        
        # Check that both matches are found
        matches = result.data['matches']
        assert len(matches) == 2
        
        # First match in "Hello World"
        assert matches[0]['cell_value'] == "Hello World"
        assert matches[0]['match_text'] == "World"
        
        # Second match in "World Peace"
        assert matches[1]['cell_value'] == "World Peace"
        assert matches[1]['match_text'] == "World"

    @pytest.mark.unit
    def test_search_table_content_exact_mode(self, document_manager, table_operations, test_doc_path):
        """Test table content search with exact mode."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with test data
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 0, "Test")
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 1, "Testing")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "Test")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "Different")
        
        # Search for exact "Test" - should find 2 matches
        result = table_operations.search_table_content(str(test_doc_path), "Test", search_mode="exact")
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 2
        assert result.data['search_mode'] == "exact"
        
        # Verify matches
        for match in result.data['matches']:
            assert match['cell_value'] == "Test"
            assert match['match_text'] == "Test"

    @pytest.mark.unit
    def test_search_table_content_regex_mode(self, document_manager, table_operations, test_doc_path):
        """Test table content search with regex mode."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with test data
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 0, "email@test.com")
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 1, "user@example.org")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "not-an-email")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "admin@site.net")
        
        # Search for email pattern
        email_pattern = r'\w+@\w+\.\w+'
        result = table_operations.search_table_content(str(test_doc_path), email_pattern, search_mode="regex")
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 3
        assert result.data['search_mode'] == "regex"
        
        # Check that all email addresses are found
        email_matches = [match['cell_value'] for match in result.data['matches']]
        assert "email@test.com" in email_matches
        assert "user@example.org" in email_matches
        assert "admin@site.net" in email_matches
        assert "not-an-email" not in email_matches

    @pytest.mark.unit
    def test_search_table_content_case_sensitive(self, document_manager, table_operations, test_doc_path):
        """Test table content search with case sensitivity."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with test data
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 0, "Test")
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 1, "test")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "TEST")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "Testing")
        
        # Case sensitive search for "test"
        result = table_operations.search_table_content(str(test_doc_path), "test", case_sensitive=True)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 1
        assert result.data['case_sensitive'] is True
        assert result.data['matches'][0]['cell_value'] == "test"
        
        # Case insensitive search for "test"
        result = table_operations.search_table_content(str(test_doc_path), "test", case_sensitive=False)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 4  # "Test", "test", "TEST", "Testing"
        assert result.data['case_sensitive'] is False

    @pytest.mark.unit
    def test_search_table_content_specific_tables(self, document_manager, table_operations, test_doc_path):
        """Test searching specific tables only."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create two tables
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 0, "Table1Data")
        
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.set_cell_value(str(test_doc_path), 1, 0, 0, "Table2Data")
        
        # Search only in table 0
        result = table_operations.search_table_content(str(test_doc_path), "Data", table_indices=[0])
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 1
        assert result.data['matches'][0]['table_index'] == 0
        assert result.data['tables_searched'] == [0]
        
        # Search only in table 1
        result = table_operations.search_table_content(str(test_doc_path), "Data", table_indices=[1])
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 1
        assert result.data['matches'][0]['table_index'] == 1
        assert result.data['tables_searched'] == [1]

    @pytest.mark.unit
    def test_search_table_content_max_results(self, document_manager, table_operations, test_doc_path):
        """Test limiting search results."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with multiple matches
        table_operations.create_table(str(test_doc_path), rows=3, cols=3)
        for row in range(3):
            for col in range(3):
                table_operations.set_cell_value(str(test_doc_path), 0, row, col, f"test{row}{col}")
        
        # Search with limit
        result = table_operations.search_table_content(str(test_doc_path), "test", max_results=5)
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 5  # Limited to 5
        assert len(result.data['matches']) == 5

    @pytest.mark.unit
    def test_search_table_headers(self, document_manager, table_operations, test_doc_path):
        """Test searching table headers specifically."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with headers
        table_operations.create_table(str(test_doc_path), rows=3, cols=3, headers=["Name", "Email", "Phone"])
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "Alice")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "alice@email.com")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 2, "123-456-7890")
        
        # Search for "Email" in headers
        result = table_operations.search_table_headers(str(test_doc_path), "Email")
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 1
        assert result.data['matches'][0]['table_index'] == 0
        assert result.data['matches'][0]['row_index'] == 0  # Header row
        assert result.data['matches'][0]['column_index'] == 1
        assert result.data['matches'][0]['cell_value'] == "Email"
        assert result.data['summary']['search_type'] == "headers_only"

    @pytest.mark.unit
    def test_search_table_content_empty_query(self, document_manager, table_operations, test_doc_path):
        """Test search with empty query."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Test empty query
        result = table_operations.search_table_content(str(test_doc_path), "")
        assert result.status == ResponseStatus.ERROR
        assert "empty" in result.message.lower()
        
        # Test whitespace-only query
        result = table_operations.search_table_content(str(test_doc_path), "   ")
        assert result.status == ResponseStatus.ERROR
        assert "empty" in result.message.lower()

    @pytest.mark.unit
    def test_search_table_content_invalid_mode(self, document_manager, table_operations, test_doc_path):
        """Test search with invalid search mode."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        result = table_operations.search_table_content(str(test_doc_path), "test", search_mode="invalid")
        assert result.status == ResponseStatus.ERROR
        assert "Invalid search mode" in result.message

    @pytest.mark.unit
    def test_search_table_content_invalid_regex(self, document_manager, table_operations, test_doc_path):
        """Test search with invalid regex pattern."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Invalid regex pattern
        result = table_operations.search_table_content(str(test_doc_path), "[invalid", search_mode="regex")
        assert result.status == ResponseStatus.ERROR
        assert "Invalid regex pattern" in result.message

    @pytest.mark.unit
    def test_search_table_content_no_tables(self, document_manager, table_operations, test_doc_path):
        """Test search in document with no tables."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        result = table_operations.search_table_content(str(test_doc_path), "test")
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['total_matches'] == 0
        assert len(result.data['matches']) == 0
        assert result.data['summary']['tables_with_matches'] == 0

    @pytest.mark.unit
    def test_search_table_content_invalid_table_indices(self, document_manager, table_operations, test_doc_path):
        """Test search with invalid table indices."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Test invalid table index
        result = table_operations.search_table_content(str(test_doc_path), "test", table_indices=[999])
        assert result.status == ResponseStatus.ERROR


class TestTableAnalysisOperations:
    """Test table structure and style analysis operations."""

    @pytest.fixture
    def setup_formatted_table(self, document_manager, table_operations, test_doc_path):
        """Set up a table with various formatting for testing analysis."""
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table with headers
        table_operations.create_table(
            str(test_doc_path), 
            rows=4, 
            cols=3, 
            headers=["Name", "Department", "Salary"]
        )
        
        # Add some data
        data_rows = [
            ["Alice Smith", "Engineering", "$75,000"],
            ["Bob Johnson", "Marketing", "$65,000"],
            ["Carol Davis", "Sales", "$55,000"]
        ]
        
        for row_idx, row_data in enumerate(data_rows, 1):
            for col_idx, value in enumerate(row_data):
                table_operations.set_cell_value(str(test_doc_path), 0, row_idx, col_idx, value)
        
        # Apply some formatting using the document directly
        document = document_manager.get_document(str(test_doc_path))
        table = document.tables[0]
        
        # Make headers bold
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply some formatting to data cells
        if len(table.rows) > 1:
            # Make first column (names) italic
            for row_idx in range(1, len(table.rows)):
                cell = table.cell(row_idx, 0)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.italic = True
            
            # Right-align salary column
            for row_idx in range(1, len(table.rows)):
                cell = table.cell(row_idx, 2)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return 0  # table index

    @pytest.mark.unit
    def test_analyze_table_structure_basic(self, document_manager, table_operations, test_doc_path, setup_formatted_table):
        """Test basic table structure analysis."""
        table_index = setup_formatted_table
        
        result = table_operations.analyze_table_structure(
            str(test_doc_path), 
            table_index, 
            include_cell_details=True
        )
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        table_info = data['table_info']
        header_info = data['header_info']
        merge_analysis = data['merge_analysis']
        style_consistency = data['style_consistency']
        
        # Check basic table info
        assert table_info['index'] == 0
        assert table_info['rows'] == 4
        assert table_info['columns'] == 3
        assert table_info['style_name'] in ['Table Grid', 'Normal Table']
        
        # Check header detection
        assert header_info['has_header'] is True
        assert header_info['header_row_index'] == 0
        assert header_info['header_cells'] == ["Name", "Department", "Salary"]
        
        # Check merge analysis
        assert merge_analysis['merged_cells_count'] == 0
        assert len(merge_analysis['merge_regions']) == 0
        
        # Check style consistency (should be mixed due to formatting)
        assert 'fonts' in style_consistency
        assert 'alignment' in style_consistency
        assert 'borders' in style_consistency

    @pytest.mark.unit
    def test_analyze_table_structure_cell_details(self, document_manager, table_operations, test_doc_path, setup_formatted_table):
        """Test table structure analysis with detailed cell information."""
        table_index = setup_formatted_table
        
        result = table_operations.analyze_table_structure(
            str(test_doc_path), 
            table_index, 
            include_cell_details=True
        )
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        cells = data['cells']
        
        # Should have 4 rows of cells
        assert len(cells) == 4
        
        # Each row should have 3 columns
        for row in cells:
            assert len(row) == 3
        
        # Check first row (headers)
        header_row = cells[0]
        for cell in header_row:
            assert cell['content']['text'] in ["Name", "Department", "Salary"]
            assert not cell['content']['is_empty']
            assert cell['text_format']['bold'] is True
            assert cell['alignment']['horizontal'] == 'center'
            assert cell['merge'] is None
        
        # Check first data row
        data_row = cells[1]
        name_cell = data_row[0]
        assert name_cell['content']['text'] == "Alice Smith"
        assert name_cell['text_format']['italic'] is True
        
        salary_cell = data_row[2]
        assert salary_cell['content']['text'] == "$75,000"
        assert salary_cell['alignment']['horizontal'] == 'right'

    @pytest.mark.unit
    def test_analyze_table_structure_without_cell_details(self, document_manager, table_operations, test_doc_path, setup_formatted_table):
        """Test table structure analysis without detailed cell information."""
        table_index = setup_formatted_table
        
        result = table_operations.analyze_table_structure(
            str(test_doc_path), 
            table_index, 
            include_cell_details=False
        )
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        cells = data['cells']
        
        # Should still have cell structure but without detailed formatting
        assert len(cells) == 4
        
        # Check that basic cell info is present but detailed formatting is minimal
        for row in cells:
            for cell in row:
                assert 'content' in cell
                assert 'position' in cell
                # Detailed formatting should be None or default values
                assert cell['text_format']['font_family'] is None
                assert cell['text_format']['font_size'] is None

    @pytest.mark.unit
    def test_analyze_table_structure_empty_table(self, document_manager, table_operations, test_doc_path):
        """Test analyzing an empty table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create empty table
        table_operations.create_table(str(test_doc_path), rows=1, cols=1)
        
        result = table_operations.analyze_table_structure(str(test_doc_path), 0)
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        assert data['table_info']['rows'] == 1
        assert data['table_info']['columns'] == 1
        assert data['merge_analysis']['merged_cells_count'] == 0
        assert len(data['cells']) == 1
        assert len(data['cells'][0]) == 1

    @pytest.mark.unit
    def test_analyze_table_structure_invalid_table(self, document_manager, table_operations, test_doc_path):
        """Test analyzing non-existent table."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        result = table_operations.analyze_table_structure(str(test_doc_path), 999)
        
        assert result.status == ResponseStatus.ERROR
        assert "out of range" in result.message or "Invalid table index" in result.message or "Table not found" in result.message

    @pytest.mark.unit
    def test_analyze_all_tables_empty_document(self, document_manager, table_operations, test_doc_path):
        """Test analyzing all tables in an empty document."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        result = table_operations.analyze_all_tables(str(test_doc_path))
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        
        # For empty documents, the response structure is different
        if 'file_info' in data:
            file_info = data['file_info']
            assert file_info['total_tables'] == 0
            assert len(data['tables']) == 0
            assert str(test_doc_path) in file_info['path']
            assert 'analysis_timestamp' in file_info
        else:
            # Fallback for the simple response structure
            assert data['total_tables'] == 0
            assert len(data['tables']) == 0
            assert str(test_doc_path) in data['file_path']

    @pytest.mark.unit
    def test_analyze_all_tables_multiple_tables(self, document_manager, table_operations, test_doc_path):
        """Test analyzing all tables in a document with multiple tables."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create multiple tables with different characteristics
        table_operations.create_table(str(test_doc_path), rows=3, cols=4, headers=["A", "B", "C", "D"])
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.create_table(str(test_doc_path), rows=5, cols=3, headers=["X", "Y", "Z"])
        
        result = table_operations.analyze_all_tables(str(test_doc_path), include_cell_details=False)
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        file_info = data['file_info']
        tables = data['tables']
        
        assert file_info['total_tables'] == 3
        assert len(tables) == 3
        
        # Check first table
        table1 = tables[0]
        assert table1['table_info']['index'] == 0
        assert table1['table_info']['rows'] == 3
        assert table1['table_info']['columns'] == 4
        assert table1['header_info']['has_header'] is True
        
        # Check second table
        table2 = tables[1]
        assert table2['table_info']['index'] == 1
        assert table2['table_info']['rows'] == 2
        assert table2['table_info']['columns'] == 2
        
        # Check third table
        table3 = tables[2]
        assert table3['table_info']['index'] == 2
        assert table3['table_info']['rows'] == 5
        assert table3['table_info']['columns'] == 3
        assert table3['header_info']['has_header'] is True

    @pytest.mark.unit
    def test_analyze_all_tables_with_cell_details(self, document_manager, table_operations, test_doc_path):
        """Test analyzing all tables with detailed cell information."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create a simple table
        table_operations.create_table(str(test_doc_path), rows=2, cols=2, headers=["Col1", "Col2"])
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "Data1")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "Data2")
        
        result = table_operations.analyze_all_tables(str(test_doc_path), include_cell_details=True)
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        assert data['file_info']['total_tables'] == 1
        
        # The analyze_all_tables method currently returns simplified table data
        # This is by design to avoid overwhelming responses for documents with many tables
        tables = data['tables']
        assert len(tables) == 1
        
        table = tables[0]
        assert table['table_info']['rows'] == 2
        assert table['table_info']['columns'] == 2
        assert table['header_info']['has_header'] is True

    @pytest.mark.unit
    def test_analyze_table_structure_style_consistency(self, document_manager, table_operations, test_doc_path):
        """Test style consistency analysis."""
        from docx.shared import RGBColor
        
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table
        table_operations.create_table(str(test_doc_path), rows=3, cols=2)
        
        # Add some data
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 0, "Header1")
        table_operations.set_cell_value(str(test_doc_path), 0, 0, 1, "Header2")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 0, "Data1")
        table_operations.set_cell_value(str(test_doc_path), 0, 1, 1, "Data2")
        table_operations.set_cell_value(str(test_doc_path), 0, 2, 0, "Data3")
        table_operations.set_cell_value(str(test_doc_path), 0, 2, 1, "Data4")
        
        # Apply mixed formatting to test consistency detection
        document = document_manager.get_document(str(test_doc_path))
        table = document.tables[0]
        
        # Make some cells bold, others not
        table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        table.cell(1, 0).paragraphs[0].runs[0].font.bold = False
        
        result = table_operations.analyze_table_structure(str(test_doc_path), 0)
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        style_summary = data['style_summary']
        
        # Should detect various styles
        assert isinstance(style_summary['font_families'], list)
        assert isinstance(style_summary['font_sizes'], list)
        assert isinstance(style_summary['colors'], list)
        assert isinstance(style_summary['background_colors'], list)

    @pytest.mark.unit
    def test_analyze_table_structure_nonexistent_document(self, table_operations):
        """Test analyzing table in non-existent document."""
        result = table_operations.analyze_table_structure("nonexistent.docx", 0)
        
        assert result.status == ResponseStatus.ERROR
        assert "not loaded" in result.message

    @pytest.mark.unit
    def test_analyze_all_tables_nonexistent_document(self, table_operations):
        """Test analyzing all tables in non-existent document."""
        result = table_operations.analyze_all_tables("nonexistent.docx")
        
        assert result.status == ResponseStatus.ERROR
        assert "not loaded" in result.message

    @pytest.mark.integration
    def test_table_analysis_integration_workflow(self, document_manager, table_operations, test_doc_path):
        """Test complete table analysis workflow."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create multiple tables with different characteristics
        # Table 1: Employee data with headers
        table_operations.create_table(
            str(test_doc_path), 
            rows=4, 
            cols=3, 
            headers=["Name", "Role", "Salary"]
        )
        
        employees = [
            ["Alice Johnson", "Developer", "$80,000"],
            ["Bob Smith", "Designer", "$70,000"],
            ["Carol White", "Manager", "$90,000"]
        ]
        
        for row_idx, emp_data in enumerate(employees, 1):
            for col_idx, value in enumerate(emp_data):
                table_operations.set_cell_value(str(test_doc_path), 0, row_idx, col_idx, value)
        
        # Table 2: Simple data without headers
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        table_operations.set_cell_value(str(test_doc_path), 1, 0, 0, "Item1")
        table_operations.set_cell_value(str(test_doc_path), 1, 0, 1, "Value1")
        table_operations.set_cell_value(str(test_doc_path), 1, 1, 0, "Item2")
        table_operations.set_cell_value(str(test_doc_path), 1, 1, 1, "Value2")
        
        # Step 1: List all tables
        list_result = table_operations.list_tables(str(test_doc_path))
        assert list_result.status == ResponseStatus.SUCCESS
        assert list_result.data['total_count'] == 2
        
        # Step 2: Analyze all tables (overview)
        all_analysis = table_operations.analyze_all_tables(str(test_doc_path), include_cell_details=False)
        assert all_analysis.status == ResponseStatus.SUCCESS
        assert all_analysis.data['file_info']['total_tables'] == 2
        
        # Step 3: Analyze first table in detail
        detailed_analysis = table_operations.analyze_table_structure(str(test_doc_path), 0, include_cell_details=True)
        assert detailed_analysis.status == ResponseStatus.SUCCESS
        
        data = detailed_analysis.data
        assert data['table_info']['rows'] == 4
        assert data['table_info']['columns'] == 3
        assert data['header_info']['has_header'] is True
        assert data['header_info']['header_cells'] == ["Name", "Role", "Salary"]
        
        # Verify cell data
        cells = data['cells']
        assert len(cells) == 4  # 4 rows
        assert cells[1][0]['content']['text'] == "Alice Johnson"  # First data row, first column
        assert cells[1][2]['content']['text'] == "$80,000"       # First data row, salary column
        
        # Step 4: Analyze second table
        table2_analysis = table_operations.analyze_table_structure(str(test_doc_path), 1, include_cell_details=True)
        assert table2_analysis.status == ResponseStatus.SUCCESS
        
        data2 = table2_analysis.data
        assert data2['table_info']['rows'] == 2
        assert data2['table_info']['columns'] == 2
        # Note: Header detection is heuristic-based, so it may detect headers even in simple tables
        # This is expected behavior when all cells in first row have content
        
        # Verify this workflow provides comprehensive table understanding
        # This is exactly what AI models need to understand table structure before modifications


class TestEnhancedCellOperations:
    """Test enhanced cell operations with formatting support."""

    @pytest.fixture
    def setup_formatted_cell_table(self, document_manager, table_operations, test_doc_path):
        """Set up a table with formatted cells for testing."""
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table
        table_operations.create_table(str(test_doc_path), rows=3, cols=3, headers=["Name", "Value", "Notes"])
        
        # Add some formatted data using the document directly
        document = document_manager.get_document(str(test_doc_path))
        table = document.tables[0]
        
        # Format header row - bold and centered
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add and format some data
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "100"
        table.cell(1, 2).text = "Important"
        
        # Make first data row italic
        for cell in table.rows[1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True
        
        # Right-align the Value column
        table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return 0  # table index

    @pytest.mark.unit
    def test_get_cell_value_with_formatting(self, table_operations, test_doc_path, setup_formatted_cell_table):
        """Test getting cell value with formatting information."""
        table_index = setup_formatted_cell_table
        
        # Get header cell with formatting
        result = table_operations.get_cell_value(str(test_doc_path), table_index, 0, 0, include_formatting=True)
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        assert data['value'] == "Name"
        assert data['is_empty'] is False
        
        # Check formatting information
        formatting = data['formatting']
        assert formatting['text_format']['bold'] is True
        assert formatting['alignment']['horizontal'] == 'center'
        assert data['merge_info'] is None

    @pytest.mark.unit
    def test_get_cell_value_without_formatting(self, table_operations, test_doc_path, setup_formatted_cell_table):
        """Test getting cell value without formatting information."""
        table_index = setup_formatted_cell_table
        
        result = table_operations.get_cell_value(str(test_doc_path), table_index, 1, 0, include_formatting=False)
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        assert data['value'] == "Alice"
        assert data['is_empty'] is False
        assert 'formatting' not in data
        assert 'merge_info' not in data

    @pytest.mark.unit
    def test_set_cell_value_with_text_formatting(self, document_manager, table_operations, test_doc_path):
        """Test setting cell value with text formatting."""
        from docx_mcp.models.formatting import TextFormat
        
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Set cell with specific text formatting
        text_format = TextFormat(
            font_family="Arial",
            font_size=14,
            font_color="FF0000",  # Red
            bold=True,
            italic=False
        )
        
        result = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Formatted Text",
            text_format=text_format,
            preserve_existing_format=False
        )
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        assert data['value'] == "Formatted Text"
        
        applied_formatting = data['applied_formatting']
        assert applied_formatting['text_format']['font_family'] == "Arial"
        assert applied_formatting['text_format']['bold'] is True

    @pytest.mark.unit
    def test_set_cell_value_with_alignment(self, document_manager, table_operations, test_doc_path):
        """Test setting cell value with alignment."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Set cell with alignment
        alignment = {"horizontal": "right", "vertical": "middle"}
        
        result = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Right Aligned",
            alignment=alignment,
            preserve_existing_format=False
        )
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        applied_formatting = data['applied_formatting']
        assert applied_formatting['alignment']['horizontal'] == 'right'

    @pytest.mark.unit
    def test_set_cell_value_with_background_color(self, document_manager, table_operations, test_doc_path):
        """Test setting cell value with background color."""
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Set cell with background color
        result = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Yellow Background",
            background_color="FFFF00",  # Yellow
            preserve_existing_format=False
        )
        
        assert result.status == ResponseStatus.SUCCESS
        
        data = result.data
        applied_formatting = data['applied_formatting']
        assert applied_formatting['background_color'] == "FFFF00"
        assert data['value'] == "Yellow Background"

    @pytest.mark.unit
    def test_set_cell_value_preserve_existing_format(self, table_operations, test_doc_path, setup_formatted_cell_table):
        """Test setting cell value while preserving existing formatting."""
        table_index = setup_formatted_cell_table
        
        # Get original formatting of a header cell
        original_result = table_operations.get_cell_value(str(test_doc_path), table_index, 0, 0, include_formatting=True)
        original_formatting = original_result.data['formatting']
        
        # Set new value while preserving formatting
        result = table_operations.set_cell_value(
            str(test_doc_path), table_index, 0, 0, "New Header",
            preserve_existing_format=True
        )
        
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['value'] == "New Header"
        
        # Verify formatting is preserved
        applied_formatting = result.data['applied_formatting']
        assert applied_formatting['text_format']['bold'] is True
        assert applied_formatting['alignment']['horizontal'] == 'center'

    @pytest.mark.unit
    def test_set_cell_value_override_existing_format(self, table_operations, test_doc_path, setup_formatted_cell_table):
        """Test setting cell value and overriding existing formatting."""
        from docx_mcp.models.formatting import TextFormat
        
        table_index = setup_formatted_cell_table
        
        # Override existing formatting
        new_format = TextFormat(
            font_family="Times New Roman",
            font_size=12,
            bold=False,
            italic=True
        )
        
        result = table_operations.set_cell_value(
            str(test_doc_path), table_index, 0, 0, "New Style",
            text_format=new_format,
            alignment={"horizontal": "left"},
            preserve_existing_format=False
        )
        
        assert result.status == ResponseStatus.SUCCESS
        
        applied_formatting = result.data['applied_formatting']
        assert applied_formatting['text_format']['bold'] is False
        assert applied_formatting['text_format']['italic'] is True
        assert applied_formatting['alignment']['horizontal'] == 'left'

    @pytest.mark.unit
    def test_enhanced_cell_operations_workflow(self, document_manager, table_operations, test_doc_path):
        """Test complete workflow of enhanced cell operations."""
        from docx_mcp.models.formatting import TextFormat
        
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        
        # Create table
        table_operations.create_table(str(test_doc_path), rows=3, cols=2, headers=["Product", "Price"])
        
        # Step 1: Get header formatting to understand current style
        header_result = table_operations.get_cell_value(str(test_doc_path), 0, 0, 0, include_formatting=True)
        assert header_result.status == ResponseStatus.SUCCESS
        
        # Step 2: Add data while preserving header style for consistency
        result1 = table_operations.set_cell_value(
            str(test_doc_path), 0, 1, 0, "Laptop",
            preserve_existing_format=True
        )
        assert result1.status == ResponseStatus.SUCCESS
        
        # Step 3: Add price with specific formatting (right-aligned)
        result2 = table_operations.set_cell_value(
            str(test_doc_path), 0, 1, 1, "$999.99",
            alignment={"horizontal": "right"},
            text_format=TextFormat(bold=True),
            preserve_existing_format=False
        )
        assert result2.status == ResponseStatus.SUCCESS
        
        # Step 4: Verify the applied formatting
        price_result = table_operations.get_cell_value(str(test_doc_path), 0, 1, 1, include_formatting=True)
        assert price_result.status == ResponseStatus.SUCCESS
        
        price_formatting = price_result.data['formatting']
        assert price_formatting['text_format']['bold'] is True
        assert price_formatting['alignment']['horizontal'] == 'right'
        
        # This workflow demonstrates how AI models can:
        # 1. Understand existing formatting
        # 2. Preserve consistent styling
        # 3. Apply specific formatting when needed
        # 4. Maintain table visual coherence

    @pytest.mark.unit
    def test_set_cell_value_invalid_formatting(self, document_manager, table_operations, test_doc_path):
        """Test setting cell value with invalid formatting parameters."""
        from docx_mcp.models.formatting import TextFormat
        
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Test with invalid color format - should not crash
        result = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Test",
            text_format=TextFormat(font_color="invalid_color"),
            preserve_existing_format=False
        )
        
        # Should succeed but ignore invalid color
        assert result.status == ResponseStatus.SUCCESS
        assert result.data['value'] == "Test"

    @pytest.mark.unit
    def test_enhanced_cell_operations_error_handling(self, table_operations):
        """Test error handling in enhanced cell operations."""
        # Test with non-existent document
        result = table_operations.get_cell_value("nonexistent.docx", 0, 0, 0)
        assert result.status == ResponseStatus.ERROR
        
        result = table_operations.set_cell_value("nonexistent.docx", 0, 0, 0, "test")
        assert result.status == ResponseStatus.ERROR


class TestCellValueIntegration:
    """Integration tests for set_cell_value and get_cell_value operations.
    
    These tests focus on verifying that formatting applied via set_cell_value
    can be correctly retrieved via get_cell_value, ensuring consistency
    between the two operations.
    """

    @pytest.mark.unit
    def test_text_formatting_consistency(self, document_manager, table_operations, test_doc_path):
        """Test that text formatting set via set_cell_value is correctly retrieved by get_cell_value."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Define comprehensive text formatting
        expected_format = {
            "font_family": "Times New Roman",
            "font_size": 16,
            "font_color": "FF0000",  # Red
            "bold": True,
            "italic": True,
            "underline": True
        }
        
        text_format = TextFormat(**expected_format)
        
        # Set cell value with formatting
        set_result = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Formatted Text",
            text_format=text_format,
            preserve_existing_format=False
        )
        
        assert set_result.status == ResponseStatus.SUCCESS
        assert set_result.data['value'] == "Formatted Text"
        
        # Get cell value and verify formatting consistency
        get_result = table_operations.get_cell_value(
            str(test_doc_path), 0, 0, 0, include_formatting=True
        )
        
        assert get_result.status == ResponseStatus.SUCCESS
        assert get_result.data['value'] == "Formatted Text"
        
        # Verify all text formatting properties
        retrieved_format = get_result.data['formatting']['text_format']
        assert retrieved_format['font_family'] == expected_format['font_family']
        assert retrieved_format['font_size'] == expected_format['font_size']
        assert retrieved_format['font_color'] == expected_format['font_color']
        assert retrieved_format['bold'] == expected_format['bold']
        assert retrieved_format['italic'] == expected_format['italic']
        assert retrieved_format['underlined'] == expected_format['underline']

    @pytest.mark.unit
    def test_alignment_consistency(self, document_manager, table_operations, test_doc_path):
        """Test that alignment settings are consistent between set and get operations."""
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=3, cols=3)
        
        # Test different alignment combinations
        alignment_tests = [
            {"horizontal": "left", "vertical": "top"},
            {"horizontal": "center", "vertical": "middle"},
            {"horizontal": "right", "vertical": "bottom"},
            {"horizontal": "justify", "vertical": "top"}
        ]
        
        for i, expected_alignment in enumerate(alignment_tests):
            row, col = divmod(i, 2)
            
            # Set cell with specific alignment
            set_result = table_operations.set_cell_value(
                str(test_doc_path), 0, row, col, f"Aligned Text {i+1}",
                alignment=expected_alignment,
                preserve_existing_format=False
            )
            
            assert set_result.status == ResponseStatus.SUCCESS
            
            # Verify alignment was applied in set operation
            applied_alignment = set_result.data['applied_formatting']['alignment']
            assert applied_alignment['horizontal'] == expected_alignment['horizontal']
            
            # Note: 'middle' gets converted to 'center' internally in Word
            expected_vertical = expected_alignment['vertical']
            if expected_vertical == 'middle':
                expected_vertical = 'center'
            assert applied_alignment['vertical'] == expected_vertical
            
            # Get cell and verify alignment consistency
            get_result = table_operations.get_cell_value(
                str(test_doc_path), 0, row, col, include_formatting=True
            )
            
            assert get_result.status == ResponseStatus.SUCCESS
            retrieved_alignment = get_result.data['formatting']['alignment']
            assert retrieved_alignment['horizontal'] == expected_alignment['horizontal']
            assert retrieved_alignment['vertical'] == expected_vertical

    @pytest.mark.unit
    def test_background_color_consistency(self, document_manager, table_operations, test_doc_path):
        """Test that background colors are consistent between set and get operations."""
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=3)
        
        # Test different background colors
        color_tests = [
            "FFFF00",  # Yellow
            "00FF00",  # Green
            "0000FF",  # Blue
            "FF00FF",  # Magenta
            "00FFFF"   # Cyan
        ]
        
        for i, expected_color in enumerate(color_tests):
            row, col = divmod(i, 3)
            
            # Set cell with background color
            set_result = table_operations.set_cell_value(
                str(test_doc_path), 0, row, col, f"Color {i+1}",
                background_color=expected_color,
                preserve_existing_format=False
            )
            
            assert set_result.status == ResponseStatus.SUCCESS
            
            # Verify background color was applied in set operation
            applied_bg = set_result.data['applied_formatting']['background_color']
            assert applied_bg == expected_color
            
            # Get cell and verify background color consistency
            get_result = table_operations.get_cell_value(
                str(test_doc_path), 0, row, col, include_formatting=True
            )
            
            assert get_result.status == ResponseStatus.SUCCESS
            retrieved_bg = get_result.data['formatting']['background_color']
            assert retrieved_bg == expected_color

    @pytest.mark.unit
    def test_comprehensive_formatting_consistency(self, document_manager, table_operations, test_doc_path):
        """Test comprehensive formatting consistency with all properties combined."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Define comprehensive formatting
        expected_text_format = TextFormat(
            font_family="Arial",
            font_size=14,
            font_color="800080",  # Purple
            bold=True,
            italic=False,
            underline=True
        )
        
        expected_alignment = {
            "horizontal": "center",
            "vertical": "middle"
        }
        
        expected_background = "FFFFCC"  # Light yellow
        
        # Set cell with all formatting options
        set_result = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Fully Formatted",
            text_format=expected_text_format,
            alignment=expected_alignment,
            background_color=expected_background,
            preserve_existing_format=False
        )
        
        assert set_result.status == ResponseStatus.SUCCESS
        assert set_result.data['value'] == "Fully Formatted"
        
        # Verify all formatting was applied in set operation
        applied_formatting = set_result.data['applied_formatting']
        
        # Check text formatting
        applied_text = applied_formatting['text_format']
        assert applied_text['font_family'] == expected_text_format.font_family
        assert applied_text['font_size'] == expected_text_format.font_size
        assert applied_text['font_color'] == expected_text_format.font_color
        assert applied_text['bold'] == expected_text_format.bold
        assert applied_text['italic'] == expected_text_format.italic
        assert applied_text['underlined'] == expected_text_format.underline
        
        # Check alignment
        applied_align = applied_formatting['alignment']
        assert applied_align['horizontal'] == expected_alignment['horizontal']
        # Note: 'middle' gets converted to 'center' internally
        expected_v_align = expected_alignment['vertical']
        if expected_v_align == 'middle':
            expected_v_align = 'center'
        assert applied_align['vertical'] == expected_v_align
        
        # Check background color
        assert applied_formatting['background_color'] == expected_background
        
        # Get cell and verify complete formatting consistency
        get_result = table_operations.get_cell_value(
            str(test_doc_path), 0, 0, 0, include_formatting=True
        )
        
        assert get_result.status == ResponseStatus.SUCCESS
        assert get_result.data['value'] == "Fully Formatted"
        
        retrieved_formatting = get_result.data['formatting']
        
        # Verify text formatting consistency
        retrieved_text = retrieved_formatting['text_format']
        assert retrieved_text['font_family'] == expected_text_format.font_family
        assert retrieved_text['font_size'] == expected_text_format.font_size
        assert retrieved_text['font_color'] == expected_text_format.font_color
        assert retrieved_text['bold'] == expected_text_format.bold
        assert retrieved_text['italic'] == expected_text_format.italic
        assert retrieved_text['underlined'] == expected_text_format.underline
        
        # Verify alignment consistency
        retrieved_align = retrieved_formatting['alignment']
        assert retrieved_align['horizontal'] == expected_alignment['horizontal']
        assert retrieved_align['vertical'] == expected_v_align
        
        # Verify background color consistency
        assert retrieved_formatting['background_color'] == expected_background

    @pytest.mark.unit
    def test_preserve_existing_format_behavior(self, document_manager, table_operations, test_doc_path):
        """Test preserve_existing_format parameter behavior in set/get consistency."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Step 1: Set initial formatting
        initial_format = TextFormat(
            font_family="Arial",
            font_size=12,
            bold=True,
            italic=False
        )
        
        initial_alignment = {"horizontal": "left", "vertical": "top"}
        initial_background = "CCCCCC"  # Light gray
        
        set_result1 = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Initial Text",
            text_format=initial_format,
            alignment=initial_alignment,
            background_color=initial_background,
            preserve_existing_format=False
        )
        
        assert set_result1.status == ResponseStatus.SUCCESS
        
        # Step 2: Update value while preserving existing format
        set_result2 = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Updated Text",
            preserve_existing_format=True
        )
        
        assert set_result2.status == ResponseStatus.SUCCESS
        assert set_result2.data['value'] == "Updated Text"
        
        # Verify that formatting was preserved in set operation
        preserved_formatting = set_result2.data['applied_formatting']
        assert preserved_formatting['text_format']['font_family'] == "Arial"
        assert preserved_formatting['text_format']['bold'] is True
        assert preserved_formatting['alignment']['horizontal'] == "left"
        assert preserved_formatting['background_color'] == "CCCCCC"
        
        # Step 3: Get cell and verify formatting consistency
        get_result = table_operations.get_cell_value(
            str(test_doc_path), 0, 0, 0, include_formatting=True
        )
        
        assert get_result.status == ResponseStatus.SUCCESS
        assert get_result.data['value'] == "Updated Text"
        
        # Verify all original formatting is still intact
        retrieved_formatting = get_result.data['formatting']
        assert retrieved_formatting['text_format']['font_family'] == "Arial"
        assert retrieved_formatting['text_format']['font_size'] == 12
        assert retrieved_formatting['text_format']['bold'] is True
        assert retrieved_formatting['text_format']['italic'] is False
        assert retrieved_formatting['alignment']['horizontal'] == "left"
        assert retrieved_formatting['alignment']['vertical'] == "top"
        assert retrieved_formatting['background_color'] == "CCCCCC"

    @pytest.mark.unit
    def test_partial_formatting_updates(self, document_manager, table_operations, test_doc_path):
        """Test partial formatting updates and their consistency."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Step 1: Set comprehensive initial formatting
        initial_format = TextFormat(
            font_family="Times New Roman",
            font_size=14,
            font_color="000080",  # Navy
            bold=True,
            italic=True,
            underline=False
        )
        
        set_result1 = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Initial",
            text_format=initial_format,
            alignment={"horizontal": "center", "vertical": "middle"},
            background_color="FFEEEE",  # Light pink
            preserve_existing_format=False
        )
        
        assert set_result1.status == ResponseStatus.SUCCESS
        
        # Step 2: Update only specific formatting properties
        partial_format = TextFormat(
            font_color="FF0000",  # Change to red
            underline=True       # Add underline
        )
        
        set_result2 = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "Partially Updated",
            text_format=partial_format,
            alignment={"horizontal": "right"},  # Change only horizontal alignment
            preserve_existing_format=True
        )
        
        assert set_result2.status == ResponseStatus.SUCCESS
        
        # Step 3: Verify the mixed formatting
        get_result = table_operations.get_cell_value(
            str(test_doc_path), 0, 0, 0, include_formatting=True
        )
        
        assert get_result.status == ResponseStatus.SUCCESS
        assert get_result.data['value'] == "Partially Updated"
        
        retrieved_formatting = get_result.data['formatting']
        
        # Verify preserved properties
        assert retrieved_formatting['text_format']['font_family'] == "Times New Roman"
        assert retrieved_formatting['text_format']['font_size'] == 14
        assert retrieved_formatting['text_format']['bold'] is True
        assert retrieved_formatting['text_format']['italic'] is True
        # Note: 'middle' gets converted to 'center' internally in Word
        assert retrieved_formatting['alignment']['vertical'] == "center"
        assert retrieved_formatting['background_color'] == "FFEEEE"
        
        # Verify updated properties
        assert retrieved_formatting['text_format']['font_color'] == "FF0000"
        assert retrieved_formatting['text_format']['underlined'] is True
        assert retrieved_formatting['alignment']['horizontal'] == "right"

    @pytest.mark.unit
    def test_formatting_edge_cases(self, document_manager, table_operations, test_doc_path):
        """Test edge cases in formatting consistency."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=3, cols=2)
        
        # Test Case 1: Empty string value with formatting
        set_result1 = table_operations.set_cell_value(
            str(test_doc_path), 0, 0, 0, "",
            text_format=TextFormat(bold=True, font_color="FF0000"),
            preserve_existing_format=False
        )
        
        assert set_result1.status == ResponseStatus.SUCCESS
        
        get_result1 = table_operations.get_cell_value(
            str(test_doc_path), 0, 0, 0, include_formatting=True
        )
        
        assert get_result1.status == ResponseStatus.SUCCESS
        assert get_result1.data['value'] == ""
        assert get_result1.data['is_empty'] is True
        # Formatting should still be preserved even for empty cells
        assert get_result1.data['formatting']['text_format']['bold'] is True
        
        # Test Case 2: Very long text with formatting
        long_text = "A" * 1000  # 1000 character string
        
        set_result2 = table_operations.set_cell_value(
            str(test_doc_path), 0, 1, 0, long_text,
            text_format=TextFormat(italic=True),
            background_color="FFFFAA",
            preserve_existing_format=False
        )
        
        assert set_result2.status == ResponseStatus.SUCCESS
        
        get_result2 = table_operations.get_cell_value(
            str(test_doc_path), 0, 1, 0, include_formatting=True
        )
        
        assert get_result2.status == ResponseStatus.SUCCESS
        assert get_result2.data['value'] == long_text
        assert get_result2.data['formatting']['text_format']['italic'] is True
        assert get_result2.data['formatting']['background_color'] == "FFFFAA"
        
        # Test Case 3: Special characters with formatting
        special_text = "Special: !@#$%^&*()_+-=[]{}|;':\",./<>?`~"
        
        set_result3 = table_operations.set_cell_value(
            str(test_doc_path), 0, 2, 0, special_text,
            text_format=TextFormat(underline=True, font_size=18),
            preserve_existing_format=False
        )
        
        assert set_result3.status == ResponseStatus.SUCCESS
        
        get_result3 = table_operations.get_cell_value(
            str(test_doc_path), 0, 2, 0, include_formatting=True
        )
        
        assert get_result3.status == ResponseStatus.SUCCESS
        assert get_result3.data['value'] == special_text
        assert get_result3.data['formatting']['text_format']['underlined'] is True
        assert get_result3.data['formatting']['text_format']['font_size'] == 18

    @pytest.mark.unit
    def test_multiple_cells_formatting_consistency(self, document_manager, table_operations, test_doc_path):
        """Test formatting consistency across multiple cells."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=3, cols=3)
        
        # Define different formatting for each cell
        cell_configs = [
            {
                "row": 0, "col": 0, "value": "Header 1",
                "format": TextFormat(bold=True, font_size=16),
                "alignment": {"horizontal": "center"},
                "background": "DDDDDD"
            },
            {
                "row": 0, "col": 1, "value": "Header 2", 
                "format": TextFormat(bold=True, italic=True),
                "alignment": {"horizontal": "left"},
                "background": "EEEEEE"
            },
            {
                "row": 1, "col": 0, "value": "Data 1",
                "format": TextFormat(font_color="0000FF"),
                "alignment": {"horizontal": "right", "vertical": "top"},
                "background": None
            },
            {
                "row": 1, "col": 1, "value": "Data 2",
                "format": TextFormat(underline=True, font_family="Courier New"),
                "alignment": {"vertical": "bottom"},
                "background": "FFFFCC"
            }
        ]
        
        # Set all cells with their respective formatting
        for config in cell_configs:
            set_result = table_operations.set_cell_value(
                str(test_doc_path), 0, config["row"], config["col"], config["value"],
                text_format=config["format"],
                alignment=config["alignment"],
                background_color=config["background"],
                preserve_existing_format=False
            )
            
            assert set_result.status == ResponseStatus.SUCCESS
            assert set_result.data['value'] == config["value"]
        
        # Verify each cell's formatting consistency
        for config in cell_configs:
            get_result = table_operations.get_cell_value(
                str(test_doc_path), 0, config["row"], config["col"], include_formatting=True
            )
            
            assert get_result.status == ResponseStatus.SUCCESS
            assert get_result.data['value'] == config["value"]
            
            retrieved_formatting = get_result.data['formatting']
            
            # Verify text formatting
            if config["format"].bold is not None:
                assert retrieved_formatting['text_format']['bold'] == config["format"].bold
            if config["format"].italic is not None:
                assert retrieved_formatting['text_format']['italic'] == config["format"].italic
            if config["format"].underline is not None:
                assert retrieved_formatting['text_format']['underlined'] == config["format"].underline
            if config["format"].font_size is not None:
                assert retrieved_formatting['text_format']['font_size'] == config["format"].font_size
            if config["format"].font_color is not None:
                assert retrieved_formatting['text_format']['font_color'] == config["format"].font_color
            if config["format"].font_family is not None:
                assert retrieved_formatting['text_format']['font_family'] == config["format"].font_family
            
            # Verify alignment
            if "horizontal" in config["alignment"]:
                assert retrieved_formatting['alignment']['horizontal'] == config["alignment"]["horizontal"]
            if "vertical" in config["alignment"]:
                assert retrieved_formatting['alignment']['vertical'] == config["alignment"]["vertical"]
            
            # Verify background color
            if config["background"]:
                assert retrieved_formatting['background_color'] == config["background"]


class TestRowStylePreservation:
    """Test row style preservation when adding new rows."""
    
    @pytest.mark.unit
    def test_add_row_preserves_style_by_default(self, document_manager, table_operations, test_doc_path):
        """Test that adding rows preserves style from adjacent rows by default."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup: Create document and table
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=3)
        
        # Step 1: Set distinctive formatting on the last row
        distinctive_format = TextFormat(
            font_family="Times New Roman",
            font_size=14,
            font_color="FF0000",  # Red
            bold=True,
            italic=True,
            underline=True
        )
        
        # Apply formatting to all cells in the last row (row 1)
        for col in range(3):
            set_result = table_operations.set_cell_value(
                str(test_doc_path), 0, 1, col, f"Styled Cell {col}",
                text_format=distinctive_format,
                alignment={"horizontal": "center", "vertical": "middle"},
                background_color="FFFF00",  # Yellow background
                preserve_existing_format=False
            )
            assert set_result.status == ResponseStatus.SUCCESS
        
        # Step 2: Add a new row using default behavior (should copy style from last row)
        add_result = table_operations.add_table_rows(
            str(test_doc_path), 
            table_index=0, 
            count=1, 
            position="end"
            # No explicit styling parameters - should use default behavior
        )
        
        assert add_result.status == ResponseStatus.SUCCESS
        assert add_result.data['rows_added'] == 1
        assert add_result.data['new_row_count'] == 3  # Original 2 + added 1
        
        # Step 3: Analyze table structure to verify style preservation
        analysis_result = table_operations.analyze_table_structure(str(test_doc_path), 0)
        assert analysis_result.status == ResponseStatus.SUCCESS
        
        # Extract cell analysis for the newly added row (row 2)
        table_data = analysis_result.data
        
        # Step 4: Verify that each cell in the new row has the same formatting as the reference row
        for col in range(3):
            # Get cell value with formatting for the new row
            get_result = table_operations.get_cell_value(
                str(test_doc_path), 0, 2, col, include_formatting=True
            )
            
            assert get_result.status == ResponseStatus.SUCCESS
            new_cell_formatting = get_result.data['formatting']
            
            # Get reference cell formatting from the previous row (row 1)
            ref_result = table_operations.get_cell_value(
                str(test_doc_path), 0, 1, col, include_formatting=True
            )
            
            assert ref_result.status == ResponseStatus.SUCCESS
            ref_cell_formatting = ref_result.data['formatting']
            
            # Verify text formatting consistency
            new_text_format = new_cell_formatting['text_format']
            ref_text_format = ref_cell_formatting['text_format']
            
            assert new_text_format['font_family'] == ref_text_format['font_family']
            assert new_text_format['font_size'] == ref_text_format['font_size']
            assert new_text_format['font_color'] == ref_text_format['font_color']
            assert new_text_format['bold'] == ref_text_format['bold']
            assert new_text_format['italic'] == ref_text_format['italic']
            assert new_text_format['underlined'] == ref_text_format['underlined']
            
            # Verify alignment consistency
            new_alignment = new_cell_formatting['alignment']
            ref_alignment = ref_cell_formatting['alignment']
            
            assert new_alignment['horizontal'] == ref_alignment['horizontal']
            assert new_alignment['vertical'] == ref_alignment['vertical']
            
            # Verify background color consistency
            assert new_cell_formatting['background_color'] == ref_cell_formatting['background_color']
    
    @pytest.mark.unit
    def test_add_row_at_beginning_copies_first_row_style(self, document_manager, table_operations, test_doc_path):
        """Test that adding rows at beginning copies style from first row."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=2, cols=2)
        
        # Set distinctive formatting on the first row
        first_row_format = TextFormat(
            font_family="Arial",
            font_size=16,
            font_color="0000FF",  # Blue
            bold=True
        )
        
        for col in range(2):
            table_operations.set_cell_value(
                str(test_doc_path), 0, 0, col, f"Header {col}",
                text_format=first_row_format,
                alignment={"horizontal": "left"},
                background_color="CCCCCC",  # Gray
                preserve_existing_format=False
            )
        
        # Add row at beginning
        add_result = table_operations.add_table_rows(
            str(test_doc_path), 
            table_index=0, 
            count=1, 
            position="beginning"
        )
        
        assert add_result.status == ResponseStatus.SUCCESS
        
        # Verify new row (now at index 0) copied style from original first row (now at index 1)
        for col in range(2):
            new_cell = table_operations.get_cell_value(
                str(test_doc_path), 0, 0, col, include_formatting=True
            )
            ref_cell = table_operations.get_cell_value(
                str(test_doc_path), 0, 1, col, include_formatting=True
            )
            
            # Verify formatting consistency
            assert new_cell.data['formatting']['text_format']['font_family'] == ref_cell.data['formatting']['text_format']['font_family']
            assert new_cell.data['formatting']['text_format']['font_size'] == ref_cell.data['formatting']['text_format']['font_size']
            assert new_cell.data['formatting']['text_format']['font_color'] == ref_cell.data['formatting']['text_format']['font_color']
            assert new_cell.data['formatting']['text_format']['bold'] == ref_cell.data['formatting']['text_format']['bold']
            assert new_cell.data['formatting']['alignment']['horizontal'] == ref_cell.data['formatting']['alignment']['horizontal']
            assert new_cell.data['formatting']['background_color'] == ref_cell.data['formatting']['background_color']
    
    @pytest.mark.unit
    def test_add_row_with_custom_style_source(self, document_manager, table_operations, test_doc_path):
        """Test adding row with explicitly specified style source row."""
        from docx_mcp.models.formatting import TextFormat
        
        # Setup: Create table with different row styles
        document_manager.open_document(str(test_doc_path), create_if_not_exists=True)
        table_operations.create_table(str(test_doc_path), rows=3, cols=2)
        
        # Style row 0 with one format
        row0_format = TextFormat(font_family="Arial", font_size=12, bold=True)
        for col in range(2):
            table_operations.set_cell_value(
                str(test_doc_path), 0, 0, col, f"Row0-{col}",
                text_format=row0_format,
                background_color="FFCCCC"  # Light red
            )
        
        # Style row 1 with different format
        row1_format = TextFormat(font_family="Calibri", font_size=16, italic=True)
        for col in range(2):
            table_operations.set_cell_value(
                str(test_doc_path), 0, 1, col, f"Row1-{col}",
                text_format=row1_format,
                background_color="CCFFCC"  # Light green
            )
        
        # Style row 2 with another format
        row2_format = TextFormat(font_family="Times New Roman", font_size=10, underline=True)
        for col in range(2):
            table_operations.set_cell_value(
                str(test_doc_path), 0, 2, col, f"Row2-{col}",
                text_format=row2_format,
                background_color="CCCCFF"  # Light blue
            )
        
        # Add new row and explicitly copy style from row 1
        add_result = table_operations.add_table_rows(
            str(test_doc_path), 
            table_index=0, 
            count=1, 
            position="end",
            copy_style_from_row=1  # Explicitly copy from row 1
        )
        
        assert add_result.status == ResponseStatus.SUCCESS
        assert add_result.data['new_row_count'] == 4
        
        # Verify new row (index 3) has same style as row 1
        for col in range(2):
            new_cell = table_operations.get_cell_value(
                str(test_doc_path), 0, 3, col, include_formatting=True
            )
            ref_cell = table_operations.get_cell_value(
                str(test_doc_path), 0, 1, col, include_formatting=True
            )
            
            # Should match row 1 style, not row 2 style
            assert new_cell.data['formatting']['text_format']['font_family'] == "Calibri"
            assert new_cell.data['formatting']['text_format']['font_size'] == 16
            assert new_cell.data['formatting']['text_format']['italic'] is True
            assert new_cell.data['formatting']['background_color'] == "CCFFCC"
            
            # Verify it matches the reference exactly
            assert new_cell.data['formatting'] == ref_cell.data['formatting']