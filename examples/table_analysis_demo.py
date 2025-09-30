#!/usr/bin/env python3
"""
Demo script showing how to use the new table structure analysis features.

This script demonstrates the comprehensive table analysis capabilities
that help AI models understand table structure and styling before
making modifications.
"""

import os
import sys
from pathlib import Path

# Add src directory to Python path
src_path = Path(__file__).parent.parent / "src"
sys.path.insert(0, str(src_path))

from docx_mcp.core.document_manager import DocumentManager
from docx_mcp.operations.tables.table_operations import TableOperations
import json


def create_sample_document():
    """Create a sample document with various table styles for testing."""
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    doc = Document()
    doc.add_heading('Table Analysis Demo Document', 0)
    
    # Table 1: Simple table with headers
    doc.add_heading('Table 1: Simple Employee Data', level=1)
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Department'
    hdr_cells[2].text = 'Salary'
    hdr_cells[3].text = 'Start Date'
    
    # Make headers bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    data_rows = [
        ['John Smith', 'Engineering', '$75,000', '2020-01-15'],
        ['Jane Doe', 'Marketing', '$65,000', '2021-03-01'],
        ['Bob Johnson', 'Sales', '$55,000', '2019-11-20']
    ]
    
    for row_data in data_rows:
        row_cells = table1.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value
    
    # Table 2: Complex table with merged cells and formatting
    doc.add_heading('Table 2: Complex Financial Report', level=1)
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    
    # Set up the complex table structure
    cells = table2.rows[0].cells
    cells[0].text = 'Financial Report Q1 2024'
    # Merge first row across all columns (this is a simplified example)
    cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
    
    # Add headers in second row
    row2_cells = table2.rows[1].cells
    row2_cells[0].text = 'Category'
    row2_cells[1].text = 'Jan'
    row2_cells[2].text = 'Feb'
    row2_cells[3].text = 'Mar'
    
    # Add data with different formatting
    data_rows2 = [
        ['Revenue', '$10,000', '$12,000', '$15,000'],
        ['Expenses', '$8,000', '$9,000', '$10,000'],
        ['Profit', '$2,000', '$3,000', '$5,000']
    ]
    
    for i, row_data in enumerate(data_rows2):
        row_cells = table2.rows[i + 2].cells
        for j, value in enumerate(row_data):
            row_cells[j].text = value
            # Add some formatting variety
            if j == 0:  # Category column
                for paragraph in row_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            elif 'Profit' in value or '$' in value:
                # Color profit numbers green
                for paragraph in row_cells[j].paragraphs:
                    for run in paragraph.runs:
                        if '$' in run.text:
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
    
    # Table 3: Simple table with different alignment
    doc.add_heading('Table 3: Product Catalog', level=1)
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    
    # Headers
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Product'
    hdr_cells3[1].text = 'Price'
    hdr_cells3[2].text = 'Stock'
    
    # Add products
    products = [
        ['Laptop', '$999.99', '25'],
        ['Mouse', '$29.99', '150'],
        ['Keyboard', '$79.99', '75']
    ]
    
    for product in products:
        row_cells = table3.add_row().cells
        for i, value in enumerate(product):
            row_cells[i].text = value
            if i == 1:  # Price column - right align
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif i == 2:  # Stock column - center align
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def analyze_single_table(table_ops, file_path, table_index):
    """Analyze a single table and display results."""
    print(f"\n{'='*60}")
    print(f"ANALYZING TABLE {table_index}")
    print(f"{'='*60}")
    
    # Get detailed analysis
    result = table_ops.analyze_table_structure(file_path, table_index, include_cell_details=True)
    
    if not result.success:
        print(f"Error analyzing table {table_index}: {result.message}")
        return
    
    data = result.data
    table_info = data['table_info']
    header_info = data['header_info']
    merge_analysis = data['merge_analysis']
    style_consistency = data['style_consistency']
    style_summary = data['style_summary']
    
    # Display basic info
    print(f"Table Index: {table_info['index']}")
    print(f"Dimensions: {table_info['rows']} rows × {table_info['columns']} columns")
    print(f"Style: {table_info['style_name'] or 'Default'}")
    
    # Header info
    print(f"\nHeader Information:")
    print(f"  Has Headers: {header_info['has_header']}")
    if header_info['has_header']:
        print(f"  Header Row: {header_info['header_row_index']}")
        print(f"  Headers: {header_info['header_cells']}")
    
    # Merge analysis
    print(f"\nMerge Analysis:")
    print(f"  Merged Cells: {merge_analysis['merged_cells_count']}")
    if merge_analysis['merge_regions']:
        print("  Merge Regions:")
        for i, region in enumerate(merge_analysis['merge_regions']):
            print(f"    {i+1}. Type: {region['type']}, "
                  f"Rows: {region['start_row']}-{region['end_row']}, "
                  f"Cols: {region['start_col']}-{region['end_col']}")
    
    # Style consistency
    print(f"\nStyle Consistency:")
    print(f"  Fonts: {'Consistent' if style_consistency['fonts'] else 'Mixed'}")
    print(f"  Alignment: {'Consistent' if style_consistency['alignment'] else 'Mixed'}")
    print(f"  Borders: {'Consistent' if style_consistency['borders'] else 'Mixed'}")
    
    # Style summary
    print(f"\nStyle Summary:")
    if style_summary['font_families']:
        print(f"  Font Families: {', '.join(style_summary['font_families'])}")
    if style_summary['font_sizes']:
        print(f"  Font Sizes: {', '.join(map(str, style_summary['font_sizes']))}")
    if style_summary['colors']:
        print(f"  Text Colors: {len(style_summary['colors'])} unique colors")
    if style_summary['background_colors']:
        print(f"  Background Colors: {len(style_summary['background_colors'])} unique colors")
    
    # Sample cell details (first few cells)
    print(f"\nSample Cell Details (first 2 rows):")
    cells = data.get('cells', [])
    for row_idx in range(min(2, len(cells))):
        print(f"  Row {row_idx}:")
        row = cells[row_idx]
        for col_idx in range(min(3, len(row))):  # Show first 3 columns
            cell = row[col_idx]
            content = cell['content']
            text_format = cell['text_format']
            alignment = cell['alignment']
            
            print(f"    Cell [{row_idx},{col_idx}]: '{content['text'][:30]}{'...' if len(content['text']) > 30 else ''}'")
            if text_format['font_family']:
                print(f"      Font: {text_format['font_family']}, Size: {text_format['font_size']}")
            if text_format['bold'] or text_format['italic']:
                styles = []
                if text_format['bold']: styles.append('Bold')
                if text_format['italic']: styles.append('Italic')
                print(f"      Style: {', '.join(styles)}")
            if alignment['horizontal'] or alignment['vertical']:
                print(f"      Alignment: H:{alignment['horizontal']}, V:{alignment['vertical']}")
            if cell['merge']:
                merge = cell['merge']
                print(f"      Merged: {merge['type']} ({merge['span_rows']}×{merge['span_cols']})")


def analyze_all_tables(table_ops, file_path):
    """Analyze all tables and display summary."""
    print(f"\n{'='*60}")
    print("ANALYZING ALL TABLES")
    print(f"{'='*60}")
    
    result = table_ops.analyze_all_tables(file_path, include_cell_details=False)
    
    if not result.success:
        print(f"Error analyzing all tables: {result.message}")
        return
    
    data = result.data
    file_info = data['file_info']
    tables = data['tables']
    
    print(f"Document: {file_info['path']}")
    print(f"Total Tables: {file_info['total_tables']}")
    print(f"Analysis Time: {file_info['analysis_timestamp']}")
    
    print(f"\nTable Summary:")
    for table in tables:
        table_info = table['table_info']
        header_info = table['header_info']
        merge_info = table['merge_analysis']
        style_info = table['style_consistency']
        
        print(f"\n  Table {table_info['index']}:")
        print(f"    Size: {table_info['rows']}×{table_info['columns']}")
        print(f"    Style: {table_info['style_name'] or 'Default'}")
        print(f"    Headers: {'Yes' if header_info['has_header'] else 'No'}")
        print(f"    Merged Cells: {merge_info['merged_cells_count']}")
        
        consistency = []
        if style_info['fonts']: consistency.append('Fonts')
        if style_info['alignment']: consistency.append('Alignment')
        if style_info['borders']: consistency.append('Borders')
        print(f"    Consistent: {', '.join(consistency) if consistency else 'None'}")


def main():
    """Main demo function."""
    print("Table Structure Analysis Demo")
    print("=" * 60)
    
    # Setup
    doc_manager = DocumentManager()
    table_ops = TableOperations(doc_manager)
    
    # Create sample document
    sample_file = "table_analysis_sample.docx"
    print(f"Creating sample document: {sample_file}")
    
    doc = create_sample_document()
    doc.save(sample_file)
    
    try:
        # Open document
        open_result = doc_manager.open_document(sample_file)
        if not open_result.success:
            print(f"Failed to open document: {open_result.message}")
            return
        
        print(f"Document opened successfully!")
        
        # List all tables first
        list_result = table_ops.list_tables(sample_file, include_summary=True)
        if list_result.success:
            tables_data = list_result.data
            print(f"\nFound {tables_data['total_count']} tables in document")
            
            # Analyze each table individually
            for i in range(tables_data['total_count']):
                analyze_single_table(table_ops, sample_file, i)
            
            # Analyze all tables together
            analyze_all_tables(table_ops, sample_file)
            
        else:
            print(f"Failed to list tables: {list_result.message}")
    
    except Exception as e:
        print(f"Error during analysis: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # Cleanup
        if os.path.exists(sample_file):
            os.remove(sample_file)
            print(f"\nCleaned up sample file: {sample_file}")
    
    print(f"\n{'='*60}")
    print("Demo completed!")
    print("The new table analysis features provide comprehensive")
    print("information about table structure and styling, helping")
    print("AI models understand and preserve existing formatting")
    print("when making modifications.")


if __name__ == "__main__":
    main()
