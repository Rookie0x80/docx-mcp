#!/usr/bin/env python3
"""
Simple example showing how to use the table analysis API.
"""

import os
import sys
from pathlib import Path

# Add src directory to Python path
src_path = Path(__file__).parent.parent / "src"
sys.path.insert(0, str(src_path))

from docx_mcp.core.document_manager import DocumentManager
from docx_mcp.operations.tables.table_operations import TableOperations
import json


def main():
    """Simple demonstration of table analysis."""
    print("Simple Table Analysis Example")
    print("=" * 50)
    
    # Setup
    doc_manager = DocumentManager()
    table_ops = TableOperations(doc_manager)
    
    # Create a simple test document
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Document', 0)
    
    # Create a simple table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Add some data
    data = [
        ['Name', 'Age', 'City'],
        ['Alice', '25', 'New York'],
        ['Bob', '30', 'London']
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            # Make header bold
            if i == 0:
                for paragraph in row.cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    # Save test document
    test_file = "simple_test.docx"
    doc.save(test_file)
    
    try:
        # Open document using the API
        result = doc_manager.open_document(test_file)
        print(f"Document opened: {result.success}")
        
        if result.success:
            # Analyze single table
            print("\n--- Single Table Analysis ---")
            table_result = table_ops.analyze_table_structure(test_file, 0, include_cell_details=True)
            
            if table_result.success:
                data = table_result.data
                print(f"Table has {data['table_info']['rows']} rows and {data['table_info']['columns']} columns")
                print(f"Headers detected: {data['header_info']['has_header']}")
                if data['header_info']['header_cells']:
                    print(f"Header cells: {data['header_info']['header_cells']}")
                
                print(f"Style consistency - Fonts: {data['style_consistency']['fonts']}")
                print(f"Merged cells: {data['merge_analysis']['merged_cells_count']}")
                
                # Show first few cells
                print("\nFirst row cell details:")
                if data['cells'] and len(data['cells']) > 0:
                    first_row = data['cells'][0]
                    for i, cell in enumerate(first_row):
                        content = cell['content']['text']
                        is_bold = cell['text_format']['bold']
                        print(f"  Cell {i}: '{content}' (Bold: {is_bold})")
            
            # Analyze all tables
            print("\n--- All Tables Analysis ---")
            all_tables_result = table_ops.analyze_all_tables(test_file, include_cell_details=False)
            
            if all_tables_result.success:
                data = all_tables_result.data
                print(f"Total tables in document: {data['file_info']['total_tables']}")
                
                for table in data['tables']:
                    table_info = table['table_info']
                    print(f"Table {table_info['index']}: {table_info['rows']}x{table_info['columns']} ({table_info['style_name']})")
    
    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"\nCleaned up: {test_file}")
    
    print("\nThis API provides comprehensive table structure analysis")
    print("including cell merging, formatting, alignment, and style information.")
    print("Perfect for AI models that need to understand and preserve table formatting!")


if __name__ == "__main__":
    main()
